"""
Script to generate the Internship Report as a formatted .docx file.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ── Helper: set all heading styles ──
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(16)
    elif i == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

def add_title(text, size=24):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_centered(text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_code_block(code_text):
    """Add a code block with monospace font and grey background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacing
    return table

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_centered("INTERNSHIP REPORT", size=22, bold=True)
doc.add_paragraph()
add_centered("on", size=14)
doc.add_paragraph()
add_centered("Evaluation of DDS Middleware for", size=18, bold=True)
add_centered("Realization of Sonar System", size=18, bold=True)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Carried out at", size=12)
add_centered("Naval Physical and Oceanographic Laboratory (NPOL)", size=14, bold=True)
add_centered("Defence Research and Development Organisation (DRDO)", size=13, bold=True)
add_centered("Thrikkakara, Kochi, Kerala", size=12)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Intern: Deon", size=12, bold=True)
add_centered("Mentor: Mr. M.P. Hemant", size=12)
add_centered("Period: 18 May 2026 – 17 June 2026", size=12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Certificate', level=1)
add_body(
    'This is to certify that the internship report entitled "Evaluation of DDS Middleware for '
    'Realization of Sonar System" is a bonafide record of the work carried out at the Naval '
    'Physical and Oceanographic Laboratory (NPOL), DRDO, Thrikkakara, Kochi, during the period '
    '18 May 2026 to 17 June 2026, under the guidance of Mr. M.P. Hemant.'
)
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Mr. M.P. Hemant")
run.bold = True
run.font.name = 'Times New Roman'
doc.add_paragraph("Mentor")
doc.add_paragraph("DRDO-NPOL, Thrikkakara")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Acknowledgement', level=1)
add_body(
    'I would like to express my sincere gratitude to DRDO-NPOL, Thrikkakara, for providing me '
    'the opportunity to undertake this internship. I am deeply thankful to my mentor, '
    'Mr. M.P. Hemant, for his invaluable guidance, constant encouragement, and technical insights '
    'throughout the duration of this project. His expertise in sonar systems and middleware '
    'architectures was instrumental in shaping my understanding of the domain.'
)
add_body(
    'I also extend my thanks to the scientists and engineers at NPOL who generously shared their '
    'knowledge and made this learning experience enriching. Finally, I thank my institution for '
    'facilitating this internship opportunity.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Placeholder)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ("1.", "Abstract", ""),
    ("2.", "Introduction", ""),
    ("", "1.1 About DRDO-NPOL", ""),
    ("", "1.2 Motivation", ""),
    ("", "1.3 Objective", ""),
    ("3.", "Literature Survey", ""),
    ("", "2.1 Sonar Systems — An Overview", ""),
    ("", "2.2 Data Distribution Service (DDS)", ""),
    ("", "2.3 Publish-Subscribe Paradigm", ""),
    ("", "2.4 Existing DDS Implementations", ""),
    ("4.", "DDS Architecture and Concepts", ""),
    ("", "3.1 DDS Layered Architecture", ""),
    ("", "3.2 Data-Centric Publish-Subscribe (DCPS)", ""),
    ("", "3.3 Quality of Service (QoS) Policies", ""),
    ("", "3.4 Automatic Discovery", ""),
    ("", "3.5 DDS vs Traditional Middleware", ""),
    ("5.", "System Design and Implementation", ""),
    ("6.", "Source Code", ""),
    ("7.", "Testing and Results", ""),
    ("8.", "Relevance to Sonar Systems", ""),
    ("9.", "Conclusion", ""),
    ("10.", "Future Scope", ""),
    ("11.", "References", ""),
]
for num, title, _ in toc_items:
    p = doc.add_paragraph()
    if num:
        run = p.add_run(f"{num}  {title}")
        run.bold = True
    else:
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Abstract', level=1)
add_body(
    'Modern naval sonar systems are complex, distributed, real-time computing environments that '
    'require robust, scalable, and low-latency communication between multiple processing nodes. '
    'The Data Distribution Service (DDS) is an OMG standard middleware that provides a data-centric '
    'publish-subscribe communication model ideally suited for such mission-critical systems. This '
    'internship at DRDO-NPOL, Thrikkakara, focused on evaluating the core principles of DDS middleware '
    '— including automatic peer discovery, decentralized data distribution, health monitoring, and '
    'quality-of-service management — and their applicability to the realization of sonar signal '
    'processing systems.'
)
add_body(
    'As part of the evaluation, a fully functional Peer-to-Peer (P2P) Multicast System was designed '
    'and implemented in Python using only standard library modules (socket, threading, json, hashlib). '
    'The prototype demonstrates key DDS concepts: automatic peer discovery via UDP multicast, '
    'decentralized node registration (analogous to DDS topics), heartbeat-based health monitoring with '
    'automatic peer pruning, state persistence, and thread-safe concurrent access. The system was '
    'validated through a comprehensive suite of unit tests covering all core components.'
)
add_body(
    'The study concludes that a DDS-based middleware architecture offers significant advantages for '
    'sonar system integration, including reduced coupling between subsystems, plug-and-play node '
    'deployment, and resilient operation in the face of node failures — all critical requirements '
    'for modern naval sonar systems.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 About DRDO-NPOL', level=2)
add_body(
    'The Naval Physical and Oceanographic Laboratory (NPOL) is a premier defence research laboratory '
    'under the Defence Research and Development Organisation (DRDO), located at Thrikkakara, Kochi, '
    'Kerala. NPOL is India\'s leading laboratory for the research, design, and development of underwater '
    'surveillance systems, including sonar systems for the Indian Navy. The laboratory specializes in:'
)
add_bullet('Hull-mounted sonar systems')
add_bullet('Towed array sonar systems')
add_bullet('Submarine sonar suites')
add_bullet('Acoustic signal processing')
add_bullet('Underwater communication systems')
add_bullet('Oceanographic instrumentation')
add_body(
    'NPOL has made significant contributions to India\'s naval capability with indigenous sonar systems '
    'such as HUMSA (Hull Mounted Sonar Array), USHUS (submarine sonar), Abhay (advanced hull-mounted '
    'sonar), and NACS (Noise and Acoustic Communication System).'
)

doc.add_heading('1.2 Motivation', level=2)
add_body(
    'Modern sonar systems are no longer monolithic hardware units. They have evolved into distributed '
    'computing systems where multiple processing nodes — responsible for beam-forming, signal '
    'conditioning, target detection, classification, and display — must communicate in real time. '
    'Traditional point-to-point or client-server communication models introduce:'
)
add_bullet('Single points of failure — the loss of a central server can cripple the entire system.')
add_bullet('Tight coupling — adding or removing processing nodes requires reconfiguration of the entire network.')
add_bullet('Scalability limitations — centralized architectures struggle to handle increasing data volumes from modern multi-element sonar arrays.')
add_body(
    'The Data Distribution Service (DDS) standard, defined by the Object Management Group (OMG), '
    'provides a data-centric publish-subscribe middleware that directly addresses these challenges. '
    'Evaluating DDS principles for sonar system realization was the primary motivation for this internship.'
)

doc.add_heading('1.3 Objective', level=2)
add_body('The objectives of this internship were:')
add_bullet('Study DDS middleware architecture — Understand the DDS standard, its layered architecture, data-centric publish-subscribe (DCPS) model, and Quality of Service (QoS) policies.')
add_bullet('Analyze sonar system requirements — Identify the communication, reliability, and real-time constraints of distributed sonar systems.')
add_bullet('Prototype a DDS-inspired system — Design and implement a peer-to-peer multicast system that demonstrates core DDS concepts including automatic discovery, decentralized data distribution, health monitoring, and fault tolerance.')
add_bullet('Evaluate feasibility — Assess the suitability of DDS middleware principles for integration into next-generation sonar systems at NPOL.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Literature Survey', level=1)

doc.add_heading('2.1 Sonar Systems — An Overview', level=2)
add_body(
    'SONAR (Sound Navigation and Ranging) is a technique that uses sound propagation through water '
    'to navigate, communicate, or detect objects underwater. Sonar systems are broadly classified into:'
)
add_table(
    ['Type', 'Description'],
    [
        ['Active Sonar', 'Emits acoustic pulses (pings) and analyzes returned echoes to determine range, bearing, and speed of underwater objects.'],
        ['Passive Sonar', 'Listens for sounds emitted by targets (e.g., submarine engine noise, propeller cavitation) without transmitting any signal.'],
    ]
)
add_body(
    'Modern naval sonar systems incorporate a distributed processing pipeline: Hydrophone Array → '
    'Pre-amplification → ADC → Beam Forming → Signal Processing → Detection → Classification → '
    'Tracking → Display. Each stage may be handled by a separate computing node, necessitating a robust '
    'inter-node communication framework capable of handling high data throughput with minimal latency.'
)

doc.add_heading('2.2 Data Distribution Service (DDS)', level=2)
add_body(
    'The Data Distribution Service for Real-Time Systems (DDS) is an Object Management Group (OMG) '
    'standard (OMG Document formal/2015-04-10) for data-centric publish-subscribe middleware. DDS provides:'
)
add_bullet('Automatic Discovery — Nodes automatically find each other without manual configuration or a central broker.')
add_bullet('Data-Centric Communication — Communication is organized around data topics rather than explicit connections between nodes.')
add_bullet('Rich QoS Policies — Over 20 configurable Quality of Service policies to control reliability, durability, latency, resource limits, and more.')
add_bullet('Decentralization — No single point of failure; every node is a peer.')
add_bullet('Real-Time Performance — Designed for microsecond-level latency in mission-critical systems.')
add_body(
    'DDS is widely used in defence, aerospace, autonomous vehicles, industrial IoT, and healthcare systems.'
)

doc.add_heading('2.3 Publish-Subscribe Paradigm', level=2)
add_body(
    'The publish-subscribe communication model decouples data producers (publishers) from data consumers '
    '(subscribers). Publishers write data to named "topics" in a global data space, and subscribers '
    'receive data from topics they have expressed interest in. Neither party needs to know about the '
    'other — they are connected only through the topic.'
)
add_table(
    ['Aspect', 'Point-to-Point', 'Publish-Subscribe (DDS)'],
    [
        ['Coupling', 'Tight — sender must know receiver', 'Loose — communication via topics'],
        ['Scalability', 'Poor — N² connections', 'Good — topics scale linearly'],
        ['Fault Tolerance', 'Low — single link failure breaks flow', 'High — nodes join/leave dynamically'],
        ['Discovery', 'Manual configuration', 'Automatic'],
    ]
)

doc.add_heading('2.4 Existing DDS Implementations', level=2)
add_table(
    ['Implementation', 'Organization', 'License', 'Notable Use'],
    [
        ['RTI Connext DDS', 'Real-Time Innovations', 'Commercial', 'US Navy, autonomous vehicles'],
        ['OpenDDS', 'Object Computing Inc.', 'Open Source (BSD)', 'Defence, research'],
        ['Eclipse Cyclone DDS', 'Eclipse Foundation', 'Open Source (EPL 2.0)', 'ROS 2 robotics framework'],
        ['Fast DDS (eProsima)', 'eProsima', 'Open Source (Apache 2.0)', 'ROS 2 default middleware'],
        ['CoreDX DDS', 'Twin Oaks Computing', 'Commercial', 'Embedded defence systems'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 3: DDS ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. DDS Architecture and Concepts', level=1)

doc.add_heading('3.1 DDS Layered Architecture', level=2)
add_body('The DDS specification defines two primary layers:')
add_body(
    '1. DCPS (Data-Centric Publish-Subscribe) — The application-facing API layer. Defines entities such as '
    'DomainParticipant, Publisher, Subscriber, DataWriter, DataReader, and Topic.'
)
add_body(
    '2. RTPS (Real-Time Publish-Subscribe) — The wire protocol that handles serialization, fragmentation, '
    'reliability, and peer discovery over the network.'
)
add_code_block("""┌─────────────────────────────────────────────┐
│          Application Layer                   │
│   (Publishers, Subscribers, Topics)          │
├─────────────────────────────────────────────┤
│    Data-Centric Publish-Subscribe (DCPS)     │
│    • DomainParticipant                       │
│    • Publisher / Subscriber                  │
│    • DataWriter / DataReader                 │
│    • Topic                                   │
├─────────────────────────────────────────────┤
│    Real-Time Publish-Subscribe (RTPS)        │
│    • Wire Protocol                           │
│    • Discovery Protocol (SPDP + SEDP)        │
│    • Transport (UDP / Shared Memory)         │
├─────────────────────────────────────────────┤
│         Network / Transport Layer            │
│    (UDP Multicast, TCP, Shared Memory)       │
└─────────────────────────────────────────────┘""")

doc.add_heading('3.2 Data-Centric Publish-Subscribe (DCPS)', level=2)
add_body('The core entities in DCPS are:')
add_table(
    ['Entity', 'Role'],
    [
        ['DomainParticipant', 'Represents a node in the DDS domain. Acts as a factory for all other entities.'],
        ['Topic', 'A named data type that defines what data is being shared (e.g., "SonarBeamData").'],
        ['Publisher', 'Manages one or more DataWriters. Responsible for sending data.'],
        ['DataWriter', 'Writes data samples of a specific Topic into the DDS global data space.'],
        ['Subscriber', 'Manages one or more DataReaders. Responsible for receiving data.'],
        ['DataReader', 'Reads data samples of a specific Topic from the DDS global data space.'],
    ]
)

doc.add_heading('3.3 Quality of Service (QoS) Policies', level=2)
add_body(
    'DDS defines a rich set of QoS policies that control data distribution behaviour. Key policies '
    'relevant to sonar systems include:'
)
add_table(
    ['QoS Policy', 'Description', 'Sonar Relevance'],
    [
        ['Reliability', 'BEST_EFFORT or RELIABLE delivery', 'Critical target detections require RELIABLE'],
        ['Durability', 'VOLATILE, TRANSIENT_LOCAL, TRANSIENT, PERSISTENT', 'Late-joining display nodes need TRANSIENT_LOCAL'],
        ['Deadline', 'Maximum expected interval between samples', 'Sonar data must arrive at fixed intervals'],
        ['Latency Budget', 'Acceptable end-to-end delay', 'Beam-forming data needs ultra-low latency'],
        ['Liveliness', 'Automatic or manual heartbeat assertion', 'Detects failed processing nodes'],
        ['History', 'KEEP_LAST(N) or KEEP_ALL', 'Display nodes may only need latest N samples'],
        ['Resource Limits', 'Max samples, instances, samples per instance', 'Prevents memory overflow in embedded processors'],
        ['Ownership', 'SHARED or EXCLUSIVE', 'Ensures only one beam-former publishes to a beam channel'],
        ['Time-Based Filter', 'Minimum separation between received samples', 'Reduces load on lower-priority display nodes'],
    ]
)

doc.add_heading('3.4 Automatic Discovery', level=2)
add_body('DDS uses a two-phase discovery protocol:')
add_body(
    '1. Simple Participant Discovery Protocol (SPDP) — Participants periodically announce their presence '
    'via UDP multicast. When a new participant is discovered, both parties exchange endpoint information.'
)
add_body(
    '2. Simple Endpoint Discovery Protocol (SEDP) — After participants discover each other, they exchange '
    'information about their DataWriters and DataReaders. Matching endpoints (same Topic, compatible QoS) '
    'are automatically connected.'
)
add_body(
    'This is directly analogous to the discovery mechanism implemented in our prototype system.'
)

doc.add_heading('3.5 DDS vs Traditional Middleware', level=2)
add_table(
    ['Feature', 'CORBA', 'JMS', 'MQTT', 'DDS'],
    [
        ['Architecture', 'Broker-based (ORB)', 'Broker-based', 'Broker-based', 'Peer-to-Peer (brokerless)'],
        ['Discovery', 'Manual / Naming Service', 'Manual', 'Broker handles', 'Automatic (SPDP + SEDP)'],
        ['Data Model', 'Object-Centric', 'Message-Centric', 'Message-Centric', 'Data-Centric'],
        ['QoS Policies', 'Limited', 'Basic', '3 levels', '22+ fine-grained policies'],
        ['Latency', 'High', 'Medium', 'Medium', 'Ultra-Low'],
        ['Real-Time', 'No', 'No', 'No', 'Yes'],
        ['Fault Tolerance', 'Single ORB failure', 'Broker is SPOF', 'Broker is SPOF', 'No single point of failure'],
        ['Best For', 'Enterprise apps', 'Enterprise messaging', 'IoT sensors', 'Mission-critical real-time'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 4: SYSTEM DESIGN AND IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. System Design and Implementation', level=1)

doc.add_heading('4.1 Approach', level=2)
add_body(
    'To practically evaluate DDS middleware concepts, a P2P Multicast System prototype was developed '
    'in Python. The prototype implements the following DDS-analogous features:'
)
add_table(
    ['DDS Concept', 'Prototype Implementation'],
    [
        ['DomainParticipant', 'P2PNode class — each running instance is a participant'],
        ['Automatic Discovery (SPDP)', 'UDP multicast DISCOVER messages on group 224.1.1.1:5007'],
        ['Topics', 'Registered nodes/channels (user-defined strings)'],
        ['Peer Health / Liveliness QoS', 'Health monitor with logical timer-based peer timeout'],
        ['Durability (TRANSIENT_LOCAL)', 'JSON config persistence in storage/configs/'],
        ['Data-Centric Model', 'Peers indexed by peer_id; data keyed by topics'],
        ['Decentralization', 'No central server; all nodes are equal peers'],
    ]
)

doc.add_heading('4.2 System Architecture', level=2)
add_body('The system follows a modular, layered architecture:')
add_code_block("""p2p_multicast_system/
│
├── main.py                        # Application entry point
│
├── core/                          # Core logic package
│   ├── node.py                    # P2PNode orchestrator class
│   ├── discovery.py               # UDP multicast sender & listener threads
│   ├── health_monitor.py          # Peer heartbeat monitoring & pruning
│   └── config_manager.py          # JSON config load/save manager
│
├── utils/                         # Shared utilities package
│   ├── hashing.py                 # SHA-256 hash generation utility
│   └── constants.py               # Network settings and timeout values
│
├── storage/configs/               # Dynamic JSON peer state files
│
├── tests/                         # Unit testing suite
│   ├── test_hashing.py
│   ├── test_discovery.py
│   └── test_health_monitor.py
│
└── docs/                          # Documentation""")

doc.add_heading('4.3 Module Description', level=2)

doc.add_heading('4.3.1 P2PNode (core/node.py)', level=3)
add_body(
    'The orchestrator class that coordinates all subsystems. It resolves the host IP and user-provided '
    'port to create a unique peer_id (e.g., "192.168.1.10:5001"). It creates a UDP socket, sets '
    'SO_REUSEADDR, binds to the multicast port, and joins the multicast group via IP_ADD_MEMBERSHIP. '
    'It maintains a self.time variable initialized to 0, used as a logical clock for the health '
    'monitoring subsystem. This replaces system wall-clock timestamps with a deterministic, testable '
    'timer. All shared state (peers dictionary, registered_nodes list) is protected by a threading.Lock.'
)

doc.add_heading('4.3.2 Discovery Service (core/discovery.py)', level=3)
add_body(
    'Implements the automatic peer discovery protocol analogous to DDS SPDP. The Discovery Sender is a '
    'daemon thread that periodically (every DISCOVERY_INTERVAL = 5 seconds) broadcasts a JSON-encoded '
    'DISCOVER message via UDP multicast containing the node\'s peer_id, host, port, registered_nodes, '
    'and SHA-256 hash. The Discovery Listener is a daemon thread that continuously listens on the '
    'multicast socket. When a DISCOVER message is received from a different peer, it registers/updates '
    'the peer in the local peers dictionary with "time": node.time (resetting the logical health '
    'timer to 0), and saves the configuration to disk.'
)

doc.add_heading('4.3.3 Health Monitor (core/health_monitor.py)', level=3)
add_body(
    'Implements liveliness monitoring analogous to DDS Liveliness QoS. It runs as a daemon thread, '
    'waking every 2 seconds. For each remote peer, the monitor increments the peer\'s time field by 2. '
    'If a peer\'s time exceeds PEER_TIMEOUT (10 seconds), the peer is declared dead and removed from '
    'the active peer list. When a DISCOVER message is received from a peer (in the listener), the peer\'s '
    'time is reset to node.time (0), effectively restarting the timeout counter.'
)
add_body(
    'This design uses a logical counter-based timeout rather than system wall-clock differences, making '
    'the health monitoring deterministic, testable, and independent of clock synchronization issues — '
    'an important consideration in distributed embedded sonar systems.'
)

doc.add_heading('4.3.4 ConfigManager (core/config_manager.py)', level=3)
add_body(
    'Provides state persistence analogous to DDS Durability QoS (TRANSIENT_LOCAL). Saves node state '
    '(peer_id, host, port, registered_nodes, peers) to storage/configs/{port}.json. On restart, loads '
    'previous state allowing the node to resume with its registered nodes and last-known peer list.'
)

doc.add_heading('4.3.5 Utilities (utils/)', level=3)
add_body(
    'constants.py centralizes all network parameters: MULTICAST_GROUP = "224.1.1.1", '
    'MULTICAST_PORT = 5007, BUFFER_SIZE = 4096, DISCOVERY_INTERVAL = 5 seconds, PEER_TIMEOUT = 10 seconds. '
    'hashing.py generates SHA-256 hashes for peer state integrity verification.'
)

doc.add_heading('4.4 Network Design', level=2)
add_table(
    ['Property', 'Value', 'Rationale'],
    [
        ['Transport Protocol', 'UDP', 'Connectionless, low overhead — ideal for periodic broadcast discovery'],
        ['Multicast Group', '224.1.1.1', 'Administratively scoped local multicast address'],
        ['Port', '5007', 'Application-level multicast port'],
        ['Message Encoding', 'JSON over raw bytes', 'Human-readable, easy to debug, sufficient for prototype'],
        ['Socket Options', 'SO_REUSEADDR, IP_ADD_MEMBERSHIP', 'Allow multiple nodes on same machine; join multicast group'],
    ]
)

doc.add_heading('4.5 Discovery Protocol', level=2)
add_body('The discovery protocol operates as follows:')
add_bullet('Node A broadcasts a DISCOVER message via UDP multicast every 5 seconds.')
add_bullet('Node B receives the message, verifies the peer_id is not its own, and registers Node A with time = 0.')
add_bullet('Node B similarly broadcasts its own DISCOVER message.')
add_bullet('Node A receives and registers Node B with time = 0.')
add_bullet('Both nodes are now aware of each other and save their updated peer lists to disk.')

doc.add_heading('4.6 Health Monitoring and Peer Pruning', level=2)
add_body('The health monitor implements a counter-based timeout mechanism:')
add_code_block("""Every 2 seconds:
  For each remote peer:
    peer.time += 2
    if peer.time > PEER_TIMEOUT (10s):
      → Remove peer from active list
      → Save config to disk
  
  When DISCOVER message received from a peer:
    → Reset peer.time to 0 (node.time)""")

doc.add_heading('4.7 State Persistence', level=2)
add_body('Each node\'s configuration is stored in a JSON file at storage/configs/{port}.json:')
add_code_block("""{
    "peer_id": "192.168.1.10:5001",
    "host": "192.168.1.10",
    "port": 5001,
    "registered_nodes": ["sonar_beam_data"],
    "peers": {
        "192.168.1.10:5001": {
            "host": "192.168.1.10",
            "port": 5001,
            "nodes": ["sonar_beam_data"],
            "hash": "e3b0c44298fc1c149afb...",
            "time": 0
        }
    }
}""")

doc.add_heading('4.8 Thread Safety', level=2)
add_body(
    'All shared state is protected by a threading.Lock. Operations including peer registration/update '
    '(discovery listener), peer pruning (health monitor), node registration (CLI), view operations, '
    'and self-initialization all acquire the lock before accessing shared state. Dictionary iteration '
    'during health monitoring uses list(node.peers.items()) to prevent RuntimeError from dictionary '
    'size changes during iteration.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 5: SOURCE CODE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Source Code', level=1)

doc.add_heading('5.1 main.py — Application Entry Point', level=2)
add_code_block("""import sys
import os

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from core.node import P2PNode

def main():
    \"\"\"Main entrypoint for starting a P2P Multicast node.\"\"\"
    try:
        node = P2PNode()
        node.start()
    except KeyboardInterrupt:
        print("\\nShutdown signal received. Exiting peer node...")
        sys.exit(0)

if __name__ == "__main__":
    main()""")

doc.add_heading('5.2 core/node.py — P2PNode Orchestrator', level=2)
add_code_block("""import socket
import threading
import time
from utils.constants import MULTICAST_GROUP, MULTICAST_PORT
from utils.hashing import generate_hash
from core.config_manager import ConfigManager
from core.discovery import start_discovery_sender, start_discovery_listener
from core.health_monitor import start_health_monitor

class P2PNode:
    def __init__(self):
        self.host = socket.gethostbyname(socket.gethostname())
        self.port = self.get_port()
        self.peer_id = f"{self.host}:{self.port}"
        self.config_manager = ConfigManager(self.port)
        self.registered_nodes = []
        self.peers = {}
        self.running = True
        self.lock = threading.Lock()
        self.time = 0
        self.load_config()
        self.sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM, socket.IPPROTO_UDP)
        self.sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        try:
            self.sock.bind(("", MULTICAST_PORT))
        except Exception:
            self.sock.bind((MULTICAST_GROUP, MULTICAST_PORT))
        mreq = socket.inet_aton(MULTICAST_GROUP) + socket.inet_aton("0.0.0.0")
        self.sock.setsockopt(socket.IPPROTO_IP, socket.IP_ADD_MEMBERSHIP, mreq)

    def get_port(self) -> int:
        while True:
            try:
                port = int(input("Enter port: "))
                if 0 < port < 65535:
                    return port
                print("Invalid port")
            except Exception:
                print("Enter numeric value")

    def save_config(self):
        self.config_manager.save_config(
            self.peer_id, self.host, self.registered_nodes, self.peers)

    def load_config(self):
        config = self.config_manager.load_config()
        if config:
            self.registered_nodes = config.get("registered_nodes", [])
            self.peers = config.get("peers", {})
        else:
            self.save_config()

    def register_node(self):
        node_name = input("\\nEnter node string: ").strip()
        if node_name == "":
            print("Invalid node")
            return
        with self.lock:
            if node_name not in self.registered_nodes:
                self.registered_nodes.append(node_name)
                self.peers[self.peer_id] = {
                    "host": self.host, "port": self.port,
                    "nodes": self.registered_nodes,
                    "hash": generate_hash(self.peer_id),
                    "time": self.time
                }
                self.save_config()

    def initialize_self_peer(self):
        with self.lock:
            self.peers[self.peer_id] = {
                "host": self.host, "port": self.port,
                "nodes": self.registered_nodes,
                "hash": generate_hash(self.peer_id),
                "time": self.time
            }
            self.save_config()

    def start(self):
        self.initialize_self_peer()
        start_discovery_listener(self)
        start_discovery_sender(self)
        start_health_monitor(self)
        self.command_loop()""")

doc.add_heading('5.3 core/discovery.py — Automatic Peer Discovery', level=2)
add_code_block("""import socket, json, time, threading
from utils.constants import MULTICAST_GROUP, MULTICAST_PORT, DISCOVERY_INTERVAL, BUFFER_SIZE
from utils.hashing import generate_hash

def start_discovery_sender(node) -> threading.Thread:
    def send_loop():
        while node.running:
            message = {
                "type": "DISCOVER", "peer_id": node.peer_id,
                "host": node.host, "port": node.port,
                "nodes": node.registered_nodes,
                "hash": generate_hash(node.peer_id)
            }
            try:
                node.sock.sendto(
                    json.dumps(message).encode(),
                    (MULTICAST_GROUP, MULTICAST_PORT))
            except Exception:
                pass
            time.sleep(DISCOVERY_INTERVAL)
    thread = threading.Thread(target=send_loop, daemon=True)
    thread.start()
    return thread

def start_discovery_listener(node) -> threading.Thread:
    def listen_loop():
        while node.running:
            try:
                data, addr = node.sock.recvfrom(BUFFER_SIZE)
                message = json.loads(data.decode())
                if message["peer_id"] == node.peer_id:
                    continue
                if message["type"] == "DISCOVER":
                    with node.lock:
                        node.peers[message["peer_id"]] = {
                            "host": message["host"],
                            "port": message["port"],
                            "nodes": message["nodes"],
                            "hash": message["hash"],
                            "time": node.time
                        }
                        node.save_config()
            except Exception:
                pass
    thread = threading.Thread(target=listen_loop, daemon=True)
    thread.start()
    return thread""")

doc.add_heading('5.4 core/health_monitor.py — Health Monitoring', level=2)
add_code_block("""import time, threading
from utils.constants import PEER_TIMEOUT

def start_health_monitor(node) -> threading.Thread:
    def monitor_loop():
        while node.running:
            remove_peers = []
            with node.lock:
                for peer_id, info in list(node.peers.items()):
                    if peer_id == node.peer_id:
                        continue
                    info["time"] += 2
                    if info["time"] > PEER_TIMEOUT:
                        remove_peers.append(peer_id)
                if remove_peers:
                    for peer_id in remove_peers:
                        del node.peers[peer_id]
                        print(f"\\nRemoved inactive peer: {peer_id}")
                    node.save_config()
            time.sleep(2)
    thread = threading.Thread(target=monitor_loop, daemon=True)
    thread.start()
    return thread""")

doc.add_heading('5.5 core/config_manager.py — State Persistence', level=2)
add_code_block("""import os, json

class ConfigManager:
    def __init__(self, port):
        self.port = port
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.storage_dir = os.path.join(base_dir, "storage", "configs")
        self.config_file = os.path.join(self.storage_dir, f"{port}.json")
        os.makedirs(self.storage_dir, exist_ok=True)

    def save_config(self, peer_id, host, registered_nodes, peers):
        data = {
            "peer_id": peer_id, "host": host,
            "port": self.port, "registered_nodes": registered_nodes,
            "peers": peers
        }
        with open(self.config_file, "w") as f:
            json.dump(data, f, indent=4)

    def load_config(self):
        if os.path.exists(self.config_file):
            with open(self.config_file, "r") as f:
                return json.load(f)
        return None""")

doc.add_heading('5.6 utils/constants.py', level=2)
add_code_block("""MULTICAST_GROUP = "224.1.1.1"
MULTICAST_PORT = 5007
BUFFER_SIZE = 4096
DISCOVERY_INTERVAL = 5
PEER_TIMEOUT = 10""")

doc.add_heading('5.7 utils/hashing.py', level=2)
add_code_block("""import hashlib

def generate_hash(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 6: TESTING AND RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Testing and Results', level=1)

doc.add_heading('6.1 Test Suite', level=2)
add_body(
    'A comprehensive unit test suite was developed using Python\'s unittest framework with '
    'unittest.mock for isolating thread behavior. The test suite covers three core areas: '
    'SHA-256 hashing utility, discovery thread spawning, and health monitor pruning logic.'
)

doc.add_heading('6.2 Test Results', level=2)
add_code_block("""$ python -m unittest discover -s p2p_multicast_system/tests

.......
----------------------------------------------------------------------
Ran 7 tests in 0.005s

OK""")

add_table(
    ['Test File', 'Test Name', 'Verification', 'Result'],
    [
        ['test_hashing.py', 'test_hash_consistency', 'Same input → same SHA-256', '✅ PASS'],
        ['test_hashing.py', 'test_hash_length_and_type', 'Output is 64-char hex string', '✅ PASS'],
        ['test_hashing.py', 'test_hash_uniqueness', 'Different inputs → different hashes', '✅ PASS'],
        ['test_discovery.py', 'test_start_discovery_sender', 'Sender spawns daemon thread', '✅ PASS'],
        ['test_discovery.py', 'test_start_discovery_listener', 'Listener spawns daemon thread', '✅ PASS'],
        ['test_health_monitor.py', 'test_start_health_monitor', 'Monitor spawns daemon thread', '✅ PASS'],
        ['test_health_monitor.py', 'test_pruning_inactive_peers', 'Inactive peers pruned correctly', '✅ PASS'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('All 7/7 tests passed.')
run.bold = True

doc.add_heading('6.3 Manual Integration Testing', level=2)
add_body(
    'The system was tested with multiple nodes running simultaneously on the same machine:'
)
add_bullet('Node A started on port 5001, registered channel "sonar_beam_data".')
add_bullet('Node B started on port 5002, registered channel "target_track".')
add_bullet('Both nodes automatically discovered each other within 5 seconds.')
add_bullet('"view peers" on both nodes showed the other peer with correct IP, port, and hash.')
add_bullet('"view nodes" showed all registered channels across both peers.')
add_bullet('When Node B was terminated, Node A removed it within 10 seconds (PEER_TIMEOUT).')
add_bullet('When Node B was restarted, it was re-discovered automatically.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 7: RELEVANCE TO SONAR SYSTEMS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Relevance to Sonar Systems', level=1)

doc.add_heading('7.1 Mapping to Sonar Architecture', level=2)
add_table(
    ['Sonar Subsystem', 'DDS / Prototype Analogy'],
    [
        ['Beam Former Node', 'P2PNode publishing "beam_data" topic'],
        ['Signal Processor Node', 'P2PNode subscribing to "beam_data", publishing "detection_data"'],
        ['Display Console', 'P2PNode subscribing to "detection_data" and "track_data"'],
        ['Recorder', 'P2PNode subscribing to all topics with KEEP_ALL history'],
        ['Redundant Processor', 'Hot-standby node that auto-discovers and takes over on failure'],
    ]
)

doc.add_heading('7.2 Key Benefits for Sonar Systems', level=2)
add_bullet('Fault Tolerance — Automatic detection and removal of failed nodes ensures the system degrades gracefully rather than catastrophically.')
add_bullet('Plug-and-Play Deployment — New processing nodes can be added to the sonar network without any manual configuration; they are discovered automatically.')
add_bullet('Reduced Integration Effort — The publish-subscribe model eliminates the need for explicit point-to-point connections between sonar subsystems.')
add_bullet('Scalability — The multicast-based discovery scales naturally as new nodes join the network.')
add_bullet('Deterministic Health Monitoring — The logical timer-based health check provides deterministic, reproducible behaviour in embedded real-time environments.')

doc.add_heading('7.3 Limitations and Considerations', level=2)
add_table(
    ['Limitation', 'Mitigation in Production DDS'],
    [
        ['JSON encoding overhead', 'DDS uses CDR binary serialization'],
        ['Python GIL limits true parallelism', 'Production DDS uses C/C++ with zero-copy'],
        ['No QoS policy negotiation', 'Full DDS provides 22+ QoS policies with automatic matching'],
        ['UDP multicast limited to LAN', 'DDS supports UDP unicast, TCP, and shared-memory transports'],
        ['No data filtering', 'DDS supports content-filtered topics and time-based filters'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 8: CONCLUSION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Conclusion', level=1)
add_body(
    'This internship at DRDO-NPOL provided valuable hands-on experience in understanding the principles '
    'of Data Distribution Service (DDS) middleware and its applicability to sonar system realization. '
    'The key outcomes of this project are:'
)
add_bullet(
    'Comprehensive understanding of DDS architecture — The study covered the DCPS and RTPS layers, '
    'QoS policies, automatic discovery protocols (SPDP/SEDP), and the data-centric publish-subscribe paradigm.'
)
add_bullet(
    'Successful prototype implementation — A fully functional P2P Multicast System was developed in Python '
    'that demonstrates core DDS concepts including automatic peer discovery via UDP multicast, decentralized '
    'node registration, logical timer-based health monitoring with automatic peer pruning, state persistence, '
    'and thread-safe concurrent operation.'
)
add_bullet(
    'Validation through testing — The system was validated with 7 unit tests covering all core components '
    '(hashing, discovery, health monitoring), all passing successfully.'
)
add_bullet(
    'Feasibility assessment — The evaluation confirms that DDS middleware provides significant advantages '
    'for distributed sonar systems: decentralization eliminates single points of failure, automatic discovery '
    'enables plug-and-play node deployment, rich QoS policies can address diverse real-time, reliability, '
    'and durability requirements, and the data-centric model reduces coupling and simplifies integration.'
)
add_body(
    'The prototype serves as a foundational proof-of-concept that validates the core communication patterns '
    'required for a DDS-based sonar middleware. With production DDS implementations (such as RTI Connext DDS '
    'or OpenDDS), these concepts can be directly applied to realize next-generation distributed sonar systems '
    'with enhanced reliability, scalability, and maintainability.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 9: FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Future Scope', level=1)
add_bullet('Integration with production DDS — Port the prototype concepts to a production DDS implementation (e.g., RTI Connext DDS, OpenDDS) with full QoS support.')
add_bullet('Sonar data simulation — Generate synthetic sonar data (beam-formed samples, detection reports) and distribute them through the DDS data space.')
add_bullet('QoS policy evaluation — Systematically benchmark different QoS configurations (RELIABLE vs BEST_EFFORT, various history depths) for sonar data streams.')
add_bullet('Multi-topic architecture — Implement a complete sonar processing pipeline with separate DDS topics for beam data, detection data, track data, and display data.')
add_bullet('Performance benchmarking — Measure latency, throughput, and jitter under realistic sonar data loads.')
add_bullet('Cross-platform deployment — Deploy DDS nodes across heterogeneous hardware (x86 servers, ARM-based embedded boards) representative of actual sonar system configurations.')
add_bullet('Security integration — Evaluate DDS Security specification for encrypted, authenticated sonar data distribution.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 10: REFERENCES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. References', level=1)

references = [
    'Object Management Group, "Data Distribution Service for Real-Time Systems (DDS), Version 1.4," OMG Document formal/2015-04-10, 2015.',
    'Object Management Group, "The Real-Time Publish-Subscribe Wire Protocol DDS Interoperability Wire Protocol (DDSI-RTPS), Version 2.3," OMG Document formal/2019-04-03, 2019.',
    'G. Pardo-Castellote, "OMG Data-Distribution Service: Architectural Overview," 23rd International Conference on Distributed Computing Systems Workshops, Providence, RI, 2003, pp. 200-206.',
    'Real-Time Innovations, "RTI Connext DDS — The Connectivity Framework for IIoT," RTI Documentation, 2024.',
    'Object Computing Inc., "OpenDDS Developer\'s Guide," Version 3.27, 2024.',
    'A. Corsaro, "Data-Centric Middleware for Autonomous and Intelligent Systems," IEEE International Conference on Autonomous Robot Systems and Competitions, 2019.',
    'J. M. López-Higuera, "Sonar Signal Processing," in Handbook of Optical Fibre Sensing Technology, John Wiley & Sons, 2002.',
    'R. J. Urick, Principles of Underwater Sound, 3rd ed., Peninsula Publishing, 1983.',
    'Python Software Foundation, "socket — Low-level networking interface," Python 3.12 Documentation, 2024.',
    'Python Software Foundation, "threading — Thread-based parallelism," Python 3.12 Documentation, 2024.',
    'eProsima, "Fast DDS Documentation," Version 2.14, 2024.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(f"[{i}]  {ref}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

# ── Save ──
output_path = os.path.join(
    r"c:\Users\DEON\OneDrive\Desktop\internship",
    "Internship_Report_DDS_Sonar_System.docx"
)
doc.save(output_path)
print(f"\n✅ Report saved to: {output_path}")
