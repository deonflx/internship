"""
Script to generate the Internship Report as a formatted .docx file.
Title: Implementing a Lightweight PTP-DDS for Building a Resilient Fault Diagnosis Information Sharing Network in Integrated Sonar Suites
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ── Helper: set all heading styles ──
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(16)
    elif i == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

def add_title(text, size=24):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_centered(text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_code_block(code_text):
    """Add a code block with monospace font and grey background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacing
    return table

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_centered("INTERNSHIP PROJECT REPORT", size=22, bold=True)
doc.add_paragraph()
add_centered("on", size=14)
doc.add_paragraph()
add_centered("Implementing a Lightweight Peer-to-Peer Data Distribution", size=16, bold=True)
add_centered("Service (PTP-DDS) for Building a Resilient Fault Diagnosis", size=16, bold=True)
add_centered("Information Sharing Network in Integrated Sonar Suites", size=16, bold=True)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Carried out at", size=12)
add_centered("Naval Physical and Oceanographic Laboratory (NPOL)", size=14, bold=True)
add_centered("Defence Research and Development Organisation (DRDO)", size=13, bold=True)
add_centered("Thrikkakara, Kochi, Kerala", size=12)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Intern: Deon", size=12, bold=True)
add_centered("Mentor: Mr. M.P. Hemant", size=12)
add_centered("Department: Sonar Systems Division", size=12)
add_centered("Period: 18 May 2026 – 17 June 2026", size=12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Certificate', level=1)
add_body(
    'This is to certify that the internship report entitled "Implementing a Lightweight Peer-to-Peer '
    'Data Distribution Service (PTP-DDS) for Building a Resilient Fault Diagnosis Information Sharing '
    'Network in Integrated Sonar Suites" is a bonafide record of the work carried out at the Naval '
    'Physical and Oceanographic Laboratory (NPOL), DRDO, Thrikkakara, Kochi, during the '
    'internship period 18 May 2026 to 17 June 2026, under the guidance of Mr. M.P. Hemant.'
)
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Mr. M.P. Hemant")
run.bold = True
run.font.name = 'Times New Roman'
doc.add_paragraph("Mentor")
doc.add_paragraph("DRDO-NPOL, Thrikkakara")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Acknowledgement', level=1)
add_body(
    'I would like to express my sincere gratitude to DRDO-NPOL for providing me '
    'the opportunity to undertake this internship. I am deeply thankful to my mentor, '
    'Mr. M.P. Hemant, for his invaluable guidance, constant encouragement, and technical insights '
    'throughout the duration of this project. His expertise in sonar systems and middleware '
    'architectures was instrumental in shaping my understanding of the domain.'
)
add_body(
    'I also extend my thanks to the scientists and engineers at NPOL who generously shared their '
    'knowledge and made this learning experience enriching. Finally, I thank my institution for '
    'facilitating this internship opportunity.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ("1.", "Abstract", True),
    ("2.", "Introduction", True),
    ("", "2.1 About DRDO-NPOL", False),
    ("", "2.2 Motivation", False),
    ("", "2.3 Objective", False),
    ("3.", "Problem Statement", True),
    ("4.", "Literature Survey", True),
    ("", "4.1 Sonar Systems — An Overview", False),
    ("", "4.2 Data Distribution Service (DDS)", False),
    ("", "4.3 Publish-Subscribe Paradigm", False),
    ("", "4.4 Existing DDS Implementations", False),
    ("5.", "DDS Architecture and Concepts", True),
    ("", "5.1 DDS Layered Architecture", False),
    ("", "5.2 Data-Centric Publish-Subscribe (DCPS)", False),
    ("", "5.3 Quality of Service (QoS) Policies", False),
    ("", "5.4 Automatic Discovery", False),
    ("", "5.5 DDS vs Traditional Middleware", False),
    ("6.", "Approaches for Dynamic Node Discovery", True),
    ("", "6.1 Centralized Registry", False),
    ("", "6.2 Gossip Protocol", False),
    ("", "6.3 Distributed Hash Table (DHT)", False),
    ("", "6.4 UDP Multicast-Based Discovery", False),
    ("", "6.5 Comparison and Justification", False),
    ("7.", "Requirement Analysis", True),
    ("", "7.1 Purpose", False),
    ("", "7.2 Scope", False),
    ("", "7.3 Functional Requirements", False),
    ("", "7.4 Non-Functional Requirements", False),
    ("", "7.5 Constraints", False),
    ("", "7.6 Use Cases", False),
    ("8.", "Design Architecture", True),
    ("", "8.1 Architectural Overview", False),
    ("", "8.2 Module Design", False),
    ("", "8.3 Data Flow Design", False),
    ("", "8.4 Network Design", False),
    ("", "8.5 Sequence Diagrams", False),
    ("", "8.6 Thread Safety Design", False),
    ("9.", "Implementation", True),
    ("", "9.1 Technology Stack", False),
    ("", "9.2 Module Implementation Details", False),
    ("10.", "Source Code", True),
    ("11.", "Fault Diagnostic System", True),
    ("", "11.1 Overview of Fault Diagnosis in Sonar Suites", False),
    ("", "11.2 Limitations of Current Monolithic FDS", False),
    ("", "11.3 PTP-DDS Approach to Fault Diagnosis", False),
    ("", "11.4 Fault Diagnosis Lifecycle", False),
    ("", "11.5 Advantages Over Monolithic FDS", False),
    ("12.", "Relevance to Sonar Systems", True),
    ("13.", "Conclusion", True),
    ("14.", "Future Scope", True),
    ("15.", "References", True),
]
for num, title, is_bold in toc_items:
    p = doc.add_paragraph()
    if num:
        run = p.add_run(f"{num}  {title}")
        run.bold = is_bold
    else:
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(title)
        run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 1: ABSTRACT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Abstract', level=1)
add_body(
    'Current generation Fault Diagnosis Systems (FDS) in integrated sonar suites are built around a '
    'single monolithic application running on one processor, connected to other software, hardware, '
    'and networking entities via the network. In such architectures, the FDS application itself becomes '
    'a single point of failure — if the FDS processor fails, fault diagnosis capability for the entire '
    'sonar suite is lost. This is operationally unacceptable for mission-critical naval systems where '
    'continuous fault monitoring is essential.'
)
add_body(
    'To address this limitation, this internship focused on designing and implementing a lightweight '
    'Peer-to-Peer Data Distribution Service (PTP-DDS) for building a resilient and easily scalable '
    'fault diagnosis information sharing network. The PTP-DDS approach distributes fault diagnosis '
    'intelligence across multiple peer nodes, eliminating the single point of failure inherent in '
    'centralized FDS architectures. Key capabilities include automatic peer discovery via UDP multicast, '
    'decentralized data distribution, heartbeat-based health monitoring with automatic fault detection '
    'and peer pruning, and state persistence for recovery after faults.'
)
add_body(
    'The scope of the work encompassed studying various designs for realising peer-to-peer networks '
    'and data distribution services — including centralized registries, gossip protocols, distributed '
    'hash tables, and UDP multicast — evaluating their suitability for building a fault diagnosis '
    'information sharing network, arriving at a design architecture, and building a prototype library. '
    'A fully functional prototype was implemented in Python using only standard library modules '
    '(socket, threading, json, hashlib) and validated through unit tests and fault injection scenarios.'
)
add_body(
    'The study concludes that the PTP-DDS approach offers significant advantages for fault diagnosis '
    'in integrated sonar suites: resilience through decentralization, easy scalability through '
    'automatic peer discovery, zero-configuration deployment, and graceful degradation under node '
    'failures — directly addressing the limitations of current monolithic FDS architectures.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 2: INTRODUCTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Introduction', level=1)

add_body(
    'Integrated sonar suites are complex, distributed, real-time computing environments comprising '
    'multiple interconnected processing nodes — beam-formers, signal processors, target detectors, '
    'display consoles, and recorders — that must operate continuously and reliably. Fault diagnosis '
    'is a critical function in such systems: the ability to automatically detect when a hardware or '
    'software component has failed, isolate the fault, and share diagnostic information across the '
    'system so that appropriate corrective action can be taken.'
)
add_body(
    'Current generation Fault Diagnosis Systems (FDS) in sonar suites are typically built around a '
    'single monolithic application running on a dedicated processor. This FDS application is connected '
    'to other software, hardware, and networking entities via the network, and it collects health status '
    'from all subsystems, performs diagnostic analysis, and reports faults. However, this centralized '
    'architecture makes the FDS itself a single point of failure — if the FDS processor crashes or '
    'becomes unreachable, the entire sonar suite loses its fault diagnosis capability.'
)
add_body(
    'To overcome this fundamental limitation, a lightweight Peer-to-Peer Data Distribution Service '
    '(PTP-DDS) is being designed to build a resilient and easily scalable fault diagnosis information '
    'sharing network. By distributing fault diagnosis intelligence across multiple peer nodes using '
    'DDS middleware principles, the system eliminates the single point of failure and enables any '
    'surviving node to continue monitoring system health.'
)

doc.add_heading('2.1 About DRDO-NPOL', level=2)
add_body(
    'The Naval Physical and Oceanographic Laboratory (NPOL) is a premier defence research laboratory '
    'under the Defence Research and Development Organisation (DRDO), located at Thrikkakara, Kochi, Kerala. '
    'NPOL is India\'s leading laboratory for the research, design, and development of underwater '
    'surveillance systems, including sonar systems for the Indian Navy. The laboratory specializes in:'
)
add_bullet('Hull-mounted sonar systems')
add_bullet('Towed array sonar systems')
add_bullet('Submarine sonar suites')
add_bullet('Acoustic signal processing')
add_bullet('Underwater communication systems')
add_bullet('Oceanographic instrumentation')
add_body(
    'NPOL has made significant contributions to India\'s naval capability with indigenous sonar systems '
    'such as HUMSA (Hull Mounted Sonar Array), USHUS (submarine sonar), Abhay (advanced hull-mounted sonar), '
    'and NACS (Noise and Acoustic Communication System).'
)

doc.add_heading('2.2 Motivation', level=2)
add_body(
    'The primary motivation for this work arises from the limitations of current centralized FDS '
    'architectures in integrated sonar suites:'
)
add_bullet('Single Point of Failure — The monolithic FDS application running on a single processor is itself a single point of failure. If the FDS processor fails, fault diagnosis capability for the entire sonar suite is lost.')
add_bullet('Tight Coupling — Adding or removing monitored subsystems requires modifications to the centralized FDS application and its network configuration.')
add_bullet('Scalability Limitations — As sonar suites grow in complexity with more processing nodes and subsystems, a single centralized FDS struggles to handle increasing volumes of health status data.')
add_bullet('Manual Configuration — Every new node or subsystem must be explicitly configured in the FDS application with its network address and monitoring parameters.')
add_body(
    'The Data Distribution Service (DDS) standard, defined by the Object Management Group (OMG), '
    'provides a data-centric publish-subscribe middleware that directly addresses these challenges. '
    'Designing a lightweight PTP-DDS to build a resilient fault diagnosis network was the primary '
    'motivation for this internship.'
)

doc.add_heading('2.3 Objective', level=2)
add_body('The objectives of this internship were:')
add_bullet('Study DDS middleware architecture — Understand the DDS standard, its layered architecture, data-centric publish-subscribe (DCPS) model, and Quality of Service (QoS) policies.')
add_bullet('Study various P2P network designs — Survey and compare multiple approaches (centralized registry, gossip protocol, DHT, UDP multicast) for realising peer-to-peer networks and data distribution services, evaluating their suitability for a fault diagnosis information sharing network.')
add_bullet('Design a PTP-DDS architecture — Arrive at a design architecture for a lightweight PTP-DDS suitable for building a resilient fault diagnosis network in integrated sonar suites.')
add_bullet('Build a prototype library — Implement the designed PTP-DDS as a prototype library demonstrating core capabilities including automatic discovery, decentralized data distribution, heartbeat-based fault detection, message passing, and fault tolerance.')
add_bullet('Validate fault diagnosis capabilities — Test and validate the system\'s ability to automatically detect, isolate, and respond to node failures through fault injection scenarios.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 3: PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Problem Statement', level=1)

add_body(
    'Current generation Fault Diagnosis Systems (FDS) in integrated sonar suites are built around a '
    'single monolithic application running on one dedicated processor. This FDS application is connected '
    'to other software modules, hardware subsystems, and networking entities via the ship-board LAN. '
    'It collects health status, performs diagnostic analysis, and reports faults to operators. However, '
    'this architecture introduces several critical challenges:'
)

doc.add_heading('3.1 Core Challenges', level=2)

add_body('Challenge 1: FDS as a Single Point of Failure')
add_body(
    'In the current architecture, the FDS application itself is a single point of failure. If the '
    'processor hosting the FDS crashes, hangs, or becomes network-unreachable, fault diagnosis '
    'capability for the entire sonar suite is lost. This is the most critical limitation — the system '
    'designed to detect faults is itself vulnerable to a single fault.'
)

add_body('Challenge 2: Static Network Configuration')
add_body(
    'Each subsystem monitored by the FDS must be manually configured with its network address and '
    'monitoring parameters in the FDS application. When subsystems are added, removed, or replaced '
    '(e.g., during hardware upgrades or field maintenance), the FDS configuration must be manually '
    'updated. This is time-consuming, error-prone, and operationally unacceptable in mission-critical '
    'scenarios.'
)

add_body('Challenge 3: Lack of Dynamic Discovery')
add_body(
    'Existing FDS implementations do not support automatic, zero-configuration discovery of subsystems. '
    'There is no mechanism for the FDS to autonomously detect the presence of new subsystems or the '
    'absence of failed ones without explicit configuration. Failed nodes may go undetected until '
    'manual inspection.'
)

add_body('Challenge 4: Scalability Constraints')
add_body(
    'As sonar suites grow in complexity with more processing nodes, sensors, and display consoles, '
    'a single centralized FDS struggles to handle increasing volumes of health status data. '
    'The centralized architecture creates a bottleneck that limits scalability.'
)

doc.add_heading('3.2 Problem Definition', level=2)
add_body(
    'The problem addressed in this project is: How can a lightweight Peer-to-Peer Data Distribution '
    'Service (PTP-DDS) be designed and implemented to build a resilient, easily scalable fault '
    'diagnosis information sharing network for integrated sonar suites — one that eliminates the '
    'single point of failure inherent in current monolithic FDS architectures, supports automatic '
    'zero-configuration peer discovery, and enables distributed fault diagnosis without any central '
    'server or manual network configuration?'
)

doc.add_heading('3.3 Scope', level=2)
add_body('The scope of this project includes:')
add_bullet('Study of various designs for realising peer-to-peer networks and data distribution services.')
add_bullet('Evaluation of their suitability for building a fault diagnosis information sharing network.')
add_bullet('Design of a PTP-DDS architecture for resilient fault diagnosis in sonar suites.')
add_bullet('Implementation of a prototype PTP-DDS library with automatic zero-configuration peer discovery using UDP multicast.')
add_bullet('Implementation of heartbeat-based health monitoring with automatic fault detection and peer pruning.')
add_bullet('Implementation of peer-to-peer directed message passing for fault information sharing.')
add_bullet('State persistence for node recovery after restarts.')
add_bullet('Unit testing, integration testing, and fault injection testing of all core modules.')

add_body('The scope excludes:')
add_bullet('Implementation of a full DDS-compliant middleware stack.')
add_bullet('Implementation of all 22+ DDS QoS policies.')
add_bullet('Real sonar data acquisition and processing.')
add_bullet('Deployment on actual sonar hardware.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 4: LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Literature Survey', level=1)

doc.add_heading('4.1 Sonar Systems — An Overview', level=2)
add_body(
    'SONAR (Sound Navigation and Ranging) is a technique that uses sound propagation through water '
    'to navigate, communicate, or detect objects underwater. Sonar systems are broadly classified into:'
)
add_table(
    ['Type', 'Description'],
    [
        ['Active Sonar', 'Emits acoustic pulses (pings) and analyzes returned echoes to determine range, bearing, and speed of underwater objects.'],
        ['Passive Sonar', 'Listens for sounds emitted by targets (e.g., submarine engine noise, propeller cavitation) without transmitting any signal.'],
    ]
)
add_body(
    'Modern naval sonar systems incorporate a distributed processing pipeline: Hydrophone Array → '
    'Pre-amplification → ADC → Beam Forming → Signal Processing → Detection → Classification → '
    'Tracking → Display. Each stage may be handled by a separate computing node, necessitating a robust '
    'inter-node communication framework capable of handling high data throughput with minimal latency.'
)

doc.add_heading('4.2 Data Distribution Service (DDS)', level=2)
add_body(
    'The Data Distribution Service for Real-Time Systems (DDS) is an Object Management Group (OMG) '
    'standard (OMG Document formal/2015-04-10) for data-centric publish-subscribe middleware. DDS provides:'
)
add_bullet('Automatic Discovery — Nodes automatically find each other without manual configuration or a central broker.')
add_bullet('Data-Centric Communication — Communication is organized around data topics rather than explicit connections between nodes.')
add_bullet('Rich QoS Policies — Over 20 configurable Quality of Service policies to control reliability, durability, latency, resource limits, and more.')
add_bullet('Decentralization — No single point of failure; every node is a peer.')
add_bullet('Real-Time Performance — Designed for microsecond-level latency in mission-critical systems.')
add_body(
    'DDS is widely used in defence, aerospace, autonomous vehicles, industrial IoT, and healthcare systems.'
)

doc.add_heading('4.3 Publish-Subscribe Paradigm', level=2)
add_body(
    'The publish-subscribe communication model decouples data producers (publishers) from data consumers '
    '(subscribers). Publishers write data to named "topics" in a global data space, and subscribers '
    'receive data from topics they have expressed interest in. Neither party needs to know about the '
    'other — they are connected only through the topic.'
)
add_table(
    ['Aspect', 'Point-to-Point', 'Publish-Subscribe (DDS)'],
    [
        ['Coupling', 'Tight — sender must know receiver', 'Loose — communication via topics'],
        ['Scalability', 'Poor — N² connections', 'Good — topics scale linearly'],
        ['Fault Tolerance', 'Low — single link failure breaks flow', 'High — nodes join/leave dynamically'],
        ['Discovery', 'Manual configuration', 'Automatic'],
    ]
)

doc.add_heading('4.4 Existing DDS Implementations', level=2)
add_table(
    ['Implementation', 'Organization', 'License', 'Notable Use'],
    [
        ['RTI Connext DDS', 'Real-Time Innovations', 'Commercial', 'US Navy, autonomous vehicles'],
        ['OpenDDS', 'Object Computing Inc.', 'Open Source (BSD)', 'Defence, research'],
        ['Eclipse Cyclone DDS', 'Eclipse Foundation', 'Open Source (EPL 2.0)', 'ROS 2 robotics framework'],
        ['Fast DDS (eProsima)', 'eProsima', 'Open Source (Apache 2.0)', 'ROS 2 default middleware'],
        ['CoreDX DDS', 'Twin Oaks Computing', 'Commercial', 'Embedded defence systems'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 5: DDS ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. DDS Architecture and Concepts', level=1)

doc.add_heading('5.1 DDS Layered Architecture', level=2)
add_body('The DDS specification defines two primary layers:')
add_body(
    '1. DCPS (Data-Centric Publish-Subscribe) — The application-facing API layer. Defines entities such as '
    'DomainParticipant, Publisher, Subscriber, DataWriter, DataReader, and Topic.'
)
add_body(
    '2. RTPS (Real-Time Publish-Subscribe) — The wire protocol that handles serialization, fragmentation, '
    'reliability, and peer discovery over the network.'
)
add_code_block("""┌─────────────────────────────────────────────┐
│          Application Layer                   │
│   (Publishers, Subscribers, Topics)          │
├─────────────────────────────────────────────┤
│    Data-Centric Publish-Subscribe (DCPS)     │
│    • DomainParticipant                       │
│    • Publisher / Subscriber                  │
│    • DataWriter / DataReader                 │
│    • Topic                                   │
├─────────────────────────────────────────────┤
│    Real-Time Publish-Subscribe (RTPS)        │
│    • Wire Protocol                           │
│    • Discovery Protocol (SPDP + SEDP)        │
│    • Transport (UDP / Shared Memory)         │
├─────────────────────────────────────────────┤
│         Network / Transport Layer            │
│    (UDP Multicast, TCP, Shared Memory)       │
└─────────────────────────────────────────────┘""")

doc.add_heading('5.2 Data-Centric Publish-Subscribe (DCPS)', level=2)
add_body('The core entities in DCPS are:')
add_table(
    ['Entity', 'Role'],
    [
        ['DomainParticipant', 'Represents a node in the DDS domain. Acts as a factory for all other entities.'],
        ['Topic', 'A named data type that defines what data is being shared (e.g., "SonarBeamData").'],
        ['Publisher', 'Manages one or more DataWriters. Responsible for sending data.'],
        ['DataWriter', 'Writes data samples of a specific Topic into the DDS global data space.'],
        ['Subscriber', 'Manages one or more DataReaders. Responsible for receiving data.'],
        ['DataReader', 'Reads data samples of a specific Topic from the DDS global data space.'],
    ]
)

doc.add_heading('5.3 Quality of Service (QoS) Policies', level=2)
add_body(
    'DDS defines a rich set of QoS policies that control data distribution behaviour. Key policies '
    'relevant to sonar systems include:'
)
add_table(
    ['QoS Policy', 'Description', 'Sonar Relevance'],
    [
        ['Reliability', 'BEST_EFFORT or RELIABLE delivery', 'Critical target detections require RELIABLE'],
        ['Durability', 'VOLATILE, TRANSIENT_LOCAL, TRANSIENT, PERSISTENT', 'Late-joining display nodes need TRANSIENT_LOCAL'],
        ['Deadline', 'Maximum expected interval between samples', 'Sonar data must arrive at fixed intervals'],
        ['Latency Budget', 'Acceptable end-to-end delay', 'Beam-forming data needs ultra-low latency'],
        ['Liveliness', 'Automatic or manual heartbeat assertion', 'Detects failed processing nodes'],
        ['History', 'KEEP_LAST(N) or KEEP_ALL', 'Display nodes may only need latest N samples'],
        ['Resource Limits', 'Max samples, instances, samples per instance', 'Prevents memory overflow in embedded processors'],
        ['Ownership', 'SHARED or EXCLUSIVE', 'Ensures only one beam-former publishes to a beam channel'],
        ['Time-Based Filter', 'Minimum separation between received samples', 'Reduces load on lower-priority display nodes'],
    ]
)

doc.add_heading('5.4 Automatic Discovery', level=2)
add_body('DDS uses a two-phase discovery protocol:')
add_body(
    '1. Simple Participant Discovery Protocol (SPDP) — Participants periodically announce their presence '
    'via UDP multicast. When a new participant is discovered, both parties exchange endpoint information.'
)
add_body(
    '2. Simple Endpoint Discovery Protocol (SEDP) — After participants discover each other, they exchange '
    'information about their DataWriters and DataReaders. Matching endpoints (same Topic, compatible QoS) '
    'are automatically connected.'
)
add_body(
    'This is directly analogous to the discovery mechanism implemented in our prototype system.'
)

doc.add_heading('5.5 DDS vs Traditional Middleware', level=2)
add_table(
    ['Feature', 'CORBA', 'JMS', 'MQTT', 'DDS'],
    [
        ['Architecture', 'Broker-based (ORB)', 'Broker-based', 'Broker-based', 'Peer-to-Peer (brokerless)'],
        ['Discovery', 'Manual / Naming Service', 'Manual', 'Broker handles', 'Automatic (SPDP + SEDP)'],
        ['Data Model', 'Object-Centric', 'Message-Centric', 'Message-Centric', 'Data-Centric'],
        ['QoS Policies', 'Limited', 'Basic', '3 levels', '22+ fine-grained policies'],
        ['Latency', 'High', 'Medium', 'Medium', 'Ultra-Low'],
        ['Real-Time', 'No', 'No', 'No', 'Yes'],
        ['Fault Tolerance', 'Single ORB failure', 'Broker is SPOF', 'Broker is SPOF', 'No single point of failure'],
        ['Best For', 'Enterprise apps', 'Enterprise messaging', 'IoT sensors', 'Mission-critical real-time'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 6: APPROACHES FOR DYNAMIC NODE DISCOVERY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Approaches for Dynamic Node Discovery in Decentralized P2P Networks', level=1)

add_body(
    'One of the most critical challenges in building a decentralized peer-to-peer system is enabling '
    'nodes to automatically discover each other without any centralized coordination or manual configuration. '
    'This chapter surveys four major approaches, analyzes their trade-offs, and justifies the approach '
    'chosen for the prototype implementation.'
)

doc.add_heading('6.1 Approach 1: Centralized Registry / Rendezvous Server', level=2)
add_body('Mechanism:')
add_body(
    'A well-known central server maintains a registry of all active nodes. When a new node joins the '
    'network, it registers with the central server, which provides it with a list of all other active nodes. '
    'Nodes periodically send heartbeats to the server to maintain their registration.'
)
add_body('Advantages:')
add_bullet('Simple to implement — straightforward client-server architecture.')
add_bullet('Fast initial discovery — a single query to the server returns the complete peer list.')
add_bullet('Works across subnets and WANs — no multicast dependency.')
add_body('Disadvantages:')
add_bullet('Single point of failure — if the registry server crashes, new nodes cannot join and existing nodes lose discovery capability.')
add_bullet('Not truly decentralized — contradicts the core P2P and DDS philosophy.')
add_bullet('Scalability bottleneck — all discovery traffic converges on a single server.')
add_bullet('Requires manual configuration — every node must be pre-configured with the server address.')
add_body('Verdict: Rejected — violates the decentralization requirement and introduces a single point of failure, '
         'which is unacceptable for mission-critical sonar systems.')

doc.add_heading('6.2 Approach 2: Gossip Protocol (Epidemic Protocol)', level=2)
add_body('Mechanism:')
add_body(
    'Each node maintains a partial view of the network. Periodically, a node randomly selects one or more '
    'known peers and exchanges its membership list. Over time, information about all nodes propagates '
    'through the network in an epidemic fashion — similar to how gossip spreads in a social group.'
)
add_body('Advantages:')
add_bullet('Fully decentralized — no central server required.')
add_bullet('Robust — tolerant of network partitions and node failures.')
add_bullet('Scalable — communication overhead grows logarithmically with network size.')
add_body('Disadvantages:')
add_bullet('Slow convergence — discovery is not immediate; it takes O(log N) rounds for information to propagate to all nodes.')
add_bullet('Eventual consistency — a newly joined node may not be known to all peers for several gossip rounds.')
add_bullet('Requires bootstrap nodes — at least one known peer must be pre-configured for initial contact (seed node problem).')
add_bullet('Complex implementation — managing partial views, merge conflicts, and membership lists adds significant complexity.')
add_bullet('Non-deterministic — propagation time depends on random peer selection.')
add_body('Verdict: Feasible but over-engineered for a LAN-scoped sonar system. The requirement for seed nodes '
         'reintroduces partial manual configuration, and the convergence delay is unacceptable for real-time systems.')

doc.add_heading('6.3 Approach 3: Distributed Hash Table (DHT) — e.g., Kademlia, Chord', level=2)
add_body('Mechanism:')
add_body(
    'Nodes are organized in a structured overlay network based on a hash function. Each node is responsible '
    'for a portion of the hash space and maintains routing tables pointing to other nodes at exponentially '
    'increasing distances. Lookups and joins follow a structured routing algorithm (e.g., Kademlia XOR-based '
    'routing achieves O(log N) hops).'
)
add_body('Advantages:')
add_bullet('Fully decentralized — no central coordinator.')
add_bullet('Deterministic lookup — guaranteed O(log N) hop lookups.')
add_bullet('Scalable — proven at massive scale (BitTorrent, IPFS).')
add_body('Disadvantages:')
add_bullet('High complexity — implementing and maintaining DHT routing tables, stabilization protocols, and finger tables is non-trivial.')
add_bullet('Requires bootstrap — at least one known node address is needed for initial join.')
add_bullet('Overhead — routing table maintenance messages add constant background traffic.')
add_bullet('Designed for WAN-scale — overkill for LAN-scoped sonar systems with tens of nodes.')
add_bullet('Slow churn handling — frequent node join/leave events (common during system maintenance) require expensive routing table updates.')
add_body('Verdict: Rejected — excessive complexity for a LAN-scoped system. DHTs are designed for internet-scale '
         'networks with thousands of nodes; a sonar system typically has 5–50 nodes on a single LAN segment.')

doc.add_heading('6.4 Approach 4: UDP Multicast-Based Discovery (Our Approach)', level=2)
add_body('Mechanism:')
add_body(
    'Each node joins a well-known IP multicast group (e.g., 224.1.1.1:5007). Periodically, every node '
    'broadcasts a DISCOVER message containing its identity (peer_id, host, port, registered nodes, hash) '
    'to the multicast group. All nodes on the LAN receive these announcements simultaneously. When a node '
    'receives a DISCOVER message from a new peer, it registers that peer in its local peer table.'
)
add_body('Advantages:')
add_bullet('Zero configuration — no seed nodes, no bootstrap servers, no manual IP addresses. A node simply joins the multicast group and is discovered.')
add_bullet('Instant discovery — within one DISCOVERY_INTERVAL (5 seconds), all nodes on the LAN are aware of each other.')
add_bullet('Truly decentralized — every node is equal; there is no coordinator or special node.')
add_bullet('No single point of failure — the discovery mechanism is distributed across all nodes.')
add_bullet('Extremely simple — the implementation requires only standard socket API calls (IP_ADD_MEMBERSHIP, sendto, recvfrom).')
add_bullet('Low overhead — a single small UDP packet per node per interval, handled efficiently by the network stack\'s multicast routing.')
add_bullet('Natural fit for LAN — multicast is natively supported on ethernet and WiFi LAN segments, which is exactly the deployment environment for sonar systems.')
add_bullet('Directly analogous to DDS SPDP — the DDS standard itself uses UDP multicast for its Simple Participant Discovery Protocol.')
add_body('Disadvantages:')
add_bullet('Limited to LAN scope — multicast is typically not routed across WAN boundaries (acceptable for sonar systems, which operate on a ship-board LAN).')
add_bullet('Multicast support required — network infrastructure must support IGMP multicast (standard on all modern LAN equipment).')
add_body('Verdict: Selected — this approach provides zero-configuration, instant, fully decentralized discovery '
         'with minimal complexity, directly mirroring the DDS SPDP protocol. It is the optimal choice for '
         'LAN-scoped sonar systems.')

doc.add_heading('6.5 Comparative Summary', level=2)
add_table(
    ['Criterion', 'Centralized Registry', 'Gossip Protocol', 'DHT (Kademlia)', 'UDP Multicast (Ours)'],
    [
        ['Decentralization', '✗ No', '✓ Yes', '✓ Yes', '✓ Yes'],
        ['Zero Configuration', '✗ No (server address)', '✗ No (seed nodes)', '✗ No (bootstrap)', '✓ Yes'],
        ['Discovery Speed', '✓ Instant', '✗ O(log N) rounds', '✗ O(log N) hops', '✓ Instant (5s)'],
        ['Single Point of Failure', '✗ Yes', '✓ None', '✓ None', '✓ None'],
        ['Implementation Complexity', '✓ Low', '✗ Medium-High', '✗ High', '✓ Very Low'],
        ['LAN Suitability', '○ OK', '○ OK', '✗ Overkill', '✓ Optimal'],
        ['WAN Suitability', '✓ Good', '✓ Good', '✓ Excellent', '✗ Not supported'],
        ['DDS Alignment', '✗ None', '○ Partial', '✗ None', '✓ Direct (SPDP)'],
        ['Fault Tolerance', '✗ Poor', '✓ Good', '✓ Good', '✓ Good'],
        ['Scalability', '✗ Server bottleneck', '✓ Logarithmic', '✓ Logarithmic', '○ LAN-bounded'],
    ]
)

add_body(
    'Our UDP multicast-based approach achieves the best trade-off for the target deployment scenario: '
    'it provides instant, zero-configuration, fully decentralized discovery with direct alignment to '
    'the DDS SPDP standard, at minimal implementation complexity. The only limitation — restriction to LAN '
    'scope — is not a constraint for ship-board sonar systems, which operate on dedicated local networks.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 7: SOFTWARE REQUIREMENTS SPECIFICATION (SRS)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Requirement Analysis', level=1)

doc.add_heading('7.1 Purpose', level=2)
add_body(
    'This chapter describes the functional and non-functional requirements for the PTP-DDS library — '
    'a lightweight peer-to-peer data distribution service designed to build a resilient fault diagnosis '
    'information sharing network for integrated sonar suites. The purpose of this system is to '
    'eliminate the single point of failure in current monolithic FDS architectures by distributing '
    'fault diagnosis capability across multiple peer nodes using DDS middleware principles.'
)

doc.add_heading('7.2 Scope', level=2)
add_body('The system shall provide the following capabilities:')
add_bullet('Automatic, zero-configuration peer discovery on a LAN.')
add_bullet('Decentralized node/topic registration and dissemination.')
add_bullet('Health monitoring with automatic detection and removal of failed peers.')
add_bullet('Directed peer-to-peer message passing.')
add_bullet('Persistent state storage for node recovery after restarts.')
add_bullet('A command-line interface for node interaction.')

doc.add_heading('7.3 Functional Requirements', level=2)

add_table(
    ['Req. ID', 'Requirement', 'Priority', 'Description'],
    [
        ['FR-01', 'Peer Discovery', 'High', 'The system shall automatically discover all peers on the LAN within one discovery interval (5 seconds) without any manual configuration.'],
        ['FR-02', 'Self-Identification', 'High', 'Each node shall generate a unique peer identifier in the format host:port upon initialization.'],
        ['FR-03', 'Node Registration', 'High', 'The system shall allow users to register named nodes/channels (analogous to DDS Topics) on their peer.'],
        ['FR-04', 'Node Dissemination', 'High', 'Registered nodes shall be broadcast to all peers as part of the discovery message.'],
        ['FR-05', 'Health Monitoring', 'High', 'The system shall monitor peer liveliness using periodic heartbeat checks and remove peers that exceed the timeout threshold (10 seconds).'],
        ['FR-06', 'Message Sending', 'Medium', 'The system shall allow a user to send a directed message to a specific peer identified by its peer_id.'],
        ['FR-07', 'Message Receiving', 'Medium', 'The system shall receive and display messages directed to the local node from any peer.'],
        ['FR-08', 'State Persistence', 'Medium', 'The system shall persist its state (registered nodes, peer list, sent/received messages) to a JSON file on disk.'],
        ['FR-09', 'State Recovery', 'Medium', 'On restart, the system shall load its previous state from the JSON config file.'],
        ['FR-10', 'View Peers', 'Low', 'The system shall display all currently active peers with their IP, port, and hash.'],
        ['FR-11', 'View Nodes', 'Low', 'The system shall display all registered nodes across all discovered peers.'],
        ['FR-12', 'View Data', 'Low', 'The system shall display all received messages with source and destination information.'],
        ['FR-13', 'Graceful Shutdown', 'Low', 'The system shall support clean shutdown via the "exit" command or Ctrl+C.'],
    ]
)

doc.add_heading('7.4 Non-Functional Requirements', level=2)

add_table(
    ['Req. ID', 'Requirement', 'Category', 'Description'],
    [
        ['NFR-01', 'Real-Time Discovery', 'Performance', 'Peer discovery latency shall not exceed 5 seconds (one DISCOVERY_INTERVAL).'],
        ['NFR-02', 'Failure Detection', 'Reliability', 'Failed peers shall be detected and removed within 10 seconds (PEER_TIMEOUT).'],
        ['NFR-03', 'Thread Safety', 'Concurrency', 'All shared state access shall be protected by mutex locks to prevent race conditions.'],
        ['NFR-04', 'Zero Configuration', 'Usability', 'The system shall require no manual IP address or peer configuration — only a port number.'],
        ['NFR-05', 'Portability', 'Portability', 'The system shall run on any platform with Python 3.8+ and standard library socket/threading support.'],
        ['NFR-06', 'Modularity', 'Maintainability', 'The system shall be organized into separate, testable modules (node, discovery, health monitor, config manager, utilities).'],
        ['NFR-07', 'Standard Libraries Only', 'Dependency', 'The core system shall use only Python standard library modules (no external dependencies for runtime).'],
    ]
)

doc.add_heading('7.5 Constraints', level=2)
add_bullet('The system operates within a single LAN segment where UDP multicast is supported.')
add_bullet('All nodes must join the same multicast group (224.1.1.1:5007).')
add_bullet('The prototype uses JSON encoding; a production system would use binary serialization (e.g., CDR).')
add_bullet('Python\'s GIL limits true CPU parallelism; this is acceptable for a prototype but not for production sonar systems requiring C/C++ performance.')

doc.add_heading('7.6 Use Cases', level=2)

add_body('Use Case 1: Start a Node')
add_table(
    ['Field', 'Description'],
    [
        ['Actor', 'User'],
        ['Precondition', 'Python 3.8+ is installed; multicast-capable network is available.'],
        ['Trigger', 'User runs "python main.py" and enters a port number.'],
        ['Main Flow', '1. System detects local IP.\n2. System creates UDP multicast socket.\n3. System loads previous config (if available).\n4. System starts discovery sender, listener, and health monitor threads.\n5. System enters command loop.'],
        ['Postcondition', 'Node is active, broadcasting DISCOVER messages, and listening for peers.'],
    ]
)

add_body('Use Case 2: Automatic Peer Discovery')
add_table(
    ['Field', 'Description'],
    [
        ['Actor', 'System (automatic)'],
        ['Precondition', 'Two or more nodes are running on the same LAN.'],
        ['Trigger', 'A DISCOVER message is received from a previously unknown peer.'],
        ['Main Flow', '1. Discovery listener receives UDP multicast message.\n2. Parses JSON payload and extracts peer_id, host, port, nodes, hash.\n3. Registers peer in local peer dictionary with time = 0.\n4. Saves updated config to disk.'],
        ['Postcondition', 'Both nodes have each other in their peer lists.'],
    ]
)

add_body('Use Case 3: Send a Message')
add_table(
    ['Field', 'Description'],
    [
        ['Actor', 'User'],
        ['Precondition', 'Node is active; at least one peer is discovered.'],
        ['Trigger', 'User enters "send" command.'],
        ['Main Flow', '1. User is prompted for destination peer_id (host:port).\n2. User is prompted for message content.\n3. System appends [data, source_peer_id, dest_peer_id] to sent list.\n4. Sent list is broadcast in next discovery message.\n5. Destination peer\'s listener receives and stores the message in recieve list.'],
        ['Postcondition', 'Message is queued for delivery and persisted on both sender and receiver.'],
    ]
)

add_body('Use Case 4: Detect and Remove a Failed Peer')
add_table(
    ['Field', 'Description'],
    [
        ['Actor', 'System (automatic)'],
        ['Precondition', 'A peer has been previously discovered.'],
        ['Trigger', 'Peer stops sending DISCOVER messages (e.g., crashed or disconnected).'],
        ['Main Flow', '1. Health monitor increments peer\'s time counter by 2 every 2 seconds.\n2. After 10 seconds without a DISCOVER message, time exceeds PEER_TIMEOUT.\n3. Peer is removed from the active peer list.\n4. Config is saved to disk.'],
        ['Postcondition', 'Failed peer is removed; system continues operating with remaining peers.'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 8: DESIGN DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Design Architecture', level=1)

doc.add_heading('8.1 Architectural Overview', level=2)
add_body(
    'The system follows a modular, layered architecture inspired by the DDS specification. The architecture '
    'separates concerns into distinct modules: orchestration (node.py), communication (discovery.py), '
    'monitoring (health_monitor.py), persistence (config_manager.py), and utilities (hashing.py, constants.py).'
)
add_code_block("""p2p_multicast_system/
│
├── main.py                        # Application entry point
│
├── core/                          # Core logic package
│   ├── __init__.py
│   ├── node.py                    # P2PNode orchestrator class
│   ├── discovery.py               # UDP multicast sender & listener threads
│   ├── health_monitor.py          # Peer heartbeat monitoring & pruning
│   └── config_manager.py          # JSON config load/save manager
│
├── utils/                         # Shared utilities package
│   ├── __init__.py
│   ├── hashing.py                 # SHA-256 hash generation utility
│   └── constants.py               # Network settings and timeout values
│
├── storage/configs/               # Dynamic JSON peer state files
│
└── tests/                         # Unit testing suite
    ├── test_hashing.py
    ├── test_discovery.py
    └── test_health_monitor.py""")

doc.add_heading('8.2 Module Design', level=2)

add_body('The system consists of the following modules with clearly defined responsibilities:')
add_table(
    ['Module', 'File', 'Responsibility', 'DDS Analogy'],
    [
        ['Orchestrator', 'core/node.py', 'Coordinates all subsystems, manages shared state, provides CLI', 'DomainParticipant'],
        ['Discovery Sender', 'core/discovery.py', 'Periodically broadcasts node identity via UDP multicast', 'SPDP Announcer'],
        ['Discovery Listener', 'core/discovery.py', 'Listens for and processes incoming discovery messages', 'SPDP Listener'],
        ['Health Monitor', 'core/health_monitor.py', 'Monitors peer liveliness, prunes dead peers', 'Liveliness QoS'],
        ['Config Manager', 'core/config_manager.py', 'Saves/loads node state to/from JSON', 'Durability QoS'],
        ['Hashing Utility', 'utils/hashing.py', 'Generates SHA-256 hashes for state integrity', 'Data Integrity'],
        ['Constants', 'utils/constants.py', 'Centralizes all configurable parameters', 'Domain Configuration'],
    ]
)

doc.add_heading('8.3 Data Flow Design', level=2)
add_body('The data flow within the system operates through the following pathways:')

add_body('Discovery Data Flow (Outbound):')
add_code_block("""P2PNode state → Discovery Sender → JSON encode → UDP multicast → LAN
  (peer_id, host, port, nodes, hash, sent messages)""")

add_body('Discovery Data Flow (Inbound):')
add_code_block("""LAN → UDP multicast → Discovery Listener → JSON decode → Update peers dict
  → Check for directed messages → Save config""")

add_body('Health Monitoring Data Flow:')
add_code_block("""Health Monitor Thread (every 2s):
  For each remote peer:
    peer.time += 2
    if peer.time > PEER_TIMEOUT:
      remove peer → save config""")

add_body('Message Passing Data Flow:')
add_code_block("""Sender Node:
  User input → append [data, src, dest] to self.sent → broadcast in DISCOVER

Receiver Node:
  Discovery Listener → check message.sent for dest == self.peer_id
  → append to self.recieve → save config""")

doc.add_heading('8.4 Network Design', level=2)
add_table(
    ['Property', 'Value', 'Rationale'],
    [
        ['Transport Protocol', 'UDP', 'Connectionless, low overhead — ideal for periodic broadcast discovery'],
        ['Multicast Group', '224.1.1.1', 'Administratively scoped local multicast address'],
        ['Port', '5007', 'Application-level multicast port (common across all nodes)'],
        ['Message Encoding', 'JSON over raw bytes', 'Human-readable, debuggable; sufficient for prototype'],
        ['Buffer Size', '4096 bytes', 'Sufficient for discovery messages with moderate peer/node counts'],
        ['Discovery Interval', '5 seconds', 'Balance between responsiveness and network overhead'],
        ['Peer Timeout', '10 seconds', 'Allows for 2 missed heartbeats before declaring a peer dead'],
        ['Socket Options', 'SO_REUSEADDR, IP_ADD_MEMBERSHIP, IP_MULTICAST_IF', 'Multiple nodes on same machine; join multicast; send on correct interface'],
    ]
)

doc.add_heading('8.5 Sequence Diagrams', level=2)

add_body('Sequence 1: Peer Discovery')
add_code_block("""Node A                              Network                            Node B
  |                                    |                                    |
  |--- DISCOVER (A's info) ---------->|--- multicast delivery ------------>|
  |                                    |                                    |
  |                                    |        [B registers A, time=0]    |
  |                                    |        [B saves config]           |
  |                                    |                                    |
  |<-- multicast delivery ------------|<--- DISCOVER (B's info) ----------|
  |                                    |                                    |
  |  [A registers B, time=0]          |                                    |
  |  [A saves config]                 |                                    |""")

add_body('Sequence 2: Peer Timeout and Pruning')
add_code_block("""Node A                Health Monitor               Peer Table
  |                        |                            |
  |  [Node B stops]        |                            |
  |                        |--- tick (2s) ------------->|  B.time = 2
  |                        |--- tick (2s) ------------->|  B.time = 4
  |                        |--- tick (2s) ------------->|  B.time = 6
  |                        |--- tick (2s) ------------->|  B.time = 8
  |                        |--- tick (2s) ------------->|  B.time = 10
  |                        |--- tick (2s) ------------->|  B.time = 12 > TIMEOUT
  |                        |                            |  → REMOVE B
  |  [Print "Removed       |                            |  → Save config
  |   inactive peer: B"]   |                            |""")

add_body('Sequence 3: Message Passing')
add_code_block("""Node A (Sender)              Network              Node B (Receiver)
  |                            |                        |
  | User: send                 |                        |
  | Input: dest=B, msg="hello" |                        |
  | → self.sent.append(...)    |                        |
  |                            |                        |
  |--- DISCOVER + sent[] ---->|--- multicast -------->  |
  |                            |                        |
  |                            |    [B checks sent[]]   |
  |                            |    [dest==B.peer_id?]  |
  |                            |    [YES → append to    |
  |                            |     B.recieve, save]   |""")

doc.add_heading('8.6 Thread Safety Design', level=2)
add_body(
    'The system uses a single threading.Lock instance (self.lock) on the P2PNode to protect all shared '
    'mutable state. The following critical sections acquire the lock before accessing shared data:'
)
add_table(
    ['Thread', 'Critical Section', 'Shared State Accessed'],
    [
        ['Discovery Listener', 'Registering/updating a peer', 'self.peers, self.recieve'],
        ['Health Monitor', 'Incrementing timers, pruning dead peers', 'self.peers'],
        ['Main Thread (CLI)', 'register_node, view_nodes, view_peers, send, view_data', 'self.registered_nodes, self.peers, self.sent, self.recieve'],
        ['Initialization', 'initialize_self_peer', 'self.peers'],
    ]
)
add_body(
    'Dictionary iteration during health monitoring uses list(node.peers.items()) to create a snapshot, '
    'preventing RuntimeError from dictionary size changes during iteration. This is a well-known Python '
    'pattern for safe concurrent dictionary modification.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 9: SYSTEM IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Implementation', level=1)

doc.add_heading('9.1 Technology Stack', level=2)
add_table(
    ['Component', 'Technology', 'Justification'],
    [
        ['Language', 'Python 3.8+', 'Rapid prototyping, rich standard library, cross-platform'],
        ['Networking', 'socket (stdlib)', 'Direct UDP/multicast support without external dependencies'],
        ['Concurrency', 'threading (stdlib)', 'Lightweight threads for concurrent I/O (sender, listener, monitor)'],
        ['Serialization', 'json (stdlib)', 'Human-readable, debuggable; sufficient for prototype payloads'],
        ['Hashing', 'hashlib (stdlib)', 'SHA-256 for peer state integrity verification'],
        ['Persistence', 'json files on disk', 'Simple, portable, human-readable state storage'],
        ['Testing', 'unittest + pytest', 'Standard Python testing framework with mocking support'],
    ]
)

doc.add_heading('9.2 Module Implementation Details', level=2)

doc.add_heading('9.2.1 P2PNode (core/node.py)', level=3)
add_body(
    'The orchestrator class that coordinates all subsystems. Key implementation details:'
)
add_bullet('IP Detection: Uses a UDP socket trick — connects to 8.8.8.8:80 (Google DNS) without sending data to determine the local LAN IP address. This works reliably across platforms and avoids hardcoding or DNS lookups.')
add_bullet('Peer ID: Formed as "{host}:{port}" — a globally unique identifier for each node instance.')
add_bullet('Socket Setup: Creates a UDP socket with IPPROTO_UDP, enables SO_REUSEADDR for port sharing, binds to the multicast port, joins the multicast group via IP_ADD_MEMBERSHIP, and sets IP_MULTICAST_IF to the local WiFi interface.')
add_bullet('Logical Clock: Maintains self.time = 0 as a logical clock for the health monitoring subsystem. This replaces system wall-clock timestamps with a deterministic, testable timer.')
add_bullet('Message Queues: self.sent and self.recieve lists store directed messages as [data, source_peer_id, dest_peer_id] tuples.')

doc.add_heading('9.2.2 Discovery Service (core/discovery.py)', level=3)
add_body(
    'Implements the automatic peer discovery protocol analogous to DDS SPDP:'
)
add_bullet('Discovery Sender: A daemon thread that broadcasts a JSON DISCOVER message every 5 seconds via UDP multicast. The message includes the node\'s peer_id, host, port, registered_nodes, hash, and sent messages.')
add_bullet('Discovery Listener: A daemon thread that continuously listens on the multicast socket. On receiving a DISCOVER message from a different peer, it: (1) checks if any messages in the sender\'s sent[] are addressed to this node and appends them to self.recieve, (2) registers/updates the peer in the local peers dictionary with time = 0.')
add_bullet('Self-Filtering: Messages from the node\'s own peer_id are silently ignored to prevent self-registration.')

doc.add_heading('9.2.3 Health Monitor (core/health_monitor.py)', level=3)
add_body(
    'Implements liveliness monitoring analogous to DDS Liveliness QoS:'
)
add_bullet('Tick Interval: Wakes every 2 seconds.')
add_bullet('Timer Increment: For each remote peer, increments peer["time"] by 2.')
add_bullet('Timeout Threshold: If peer["time"] > PEER_TIMEOUT (10 seconds), the peer is declared dead and removed.')
add_bullet('Reset Mechanism: When a DISCOVER message is received (in the listener), the peer\'s time is reset to 0, effectively restarting the timeout counter.')
add_bullet('Deterministic Design: Uses a logical counter rather than wall-clock differences, making behaviour reproducible and testable.')

doc.add_heading('9.2.4 ConfigManager (core/config_manager.py)', level=3)
add_body(
    'Provides state persistence analogous to DDS Durability QoS (TRANSIENT_LOCAL):'
)
add_bullet('File Location: storage/configs/{port}.json — unique per node instance.')
add_bullet('Auto-Directory Creation: Uses os.makedirs(exist_ok=True) to ensure the storage directory exists.')
add_bullet('Save: Serializes peer_id, host, port, registered_nodes, peers, sent, and recieve to JSON with 4-space indentation.')
add_bullet('Load: Reads and parses the JSON file on startup; returns None if no file exists (first run).')

doc.add_heading('9.2.5 Utilities', level=3)
add_bullet('constants.py: Centralizes MULTICAST_GROUP (224.1.1.1), MULTICAST_PORT (5007), BUFFER_SIZE (4096), DISCOVERY_INTERVAL (5s), PEER_TIMEOUT (10s).')
add_bullet('hashing.py: Generates SHA-256 hex digest of peer_id for state integrity verification.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 10: SOURCE CODE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Source Code', level=1)

doc.add_heading('10.1 main.py — Application Entry Point', level=2)
add_code_block("""import sys
import os

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from core.node import P2PNode

def main():
    \"\"\"Main entrypoint for starting a P2P Multicast node.\"\"\"
    try:
        node = P2PNode()
        node.start()
    except KeyboardInterrupt:
        print("\\nShutdown signal received. Exiting peer node...")
        sys.exit(0)

if __name__ == "__main__":
    main()""")

doc.add_heading('10.2 core/node.py — P2PNode Orchestrator', level=2)
add_code_block("""import socket
import threading
import time
from utils.constants import MULTICAST_GROUP, MULTICAST_PORT
from utils.hashing import generate_hash
from core.config_manager import ConfigManager
from core.discovery import start_discovery_sender, start_discovery_listener
from core.health_monitor import start_health_monitor

def get_local_ip():
    \"\"\"Auto-detect the local LAN IP address.\"\"\"
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"

lip = get_local_ip()

class P2PNode:
    def __init__(self):
        self.host = lip
        self.port = self.get_port()
        self.peer_id = f"{self.host}:{self.port}"
        self.config_manager = ConfigManager(self.port)
        self.registered_nodes = []
        self.peers = {}
        self.running = True
        self.lock = threading.Lock()
        self.time = 0
        self.sent = []
        self.recieve = []
        self.load_config()

        # Create UDP Multicast socket
        self.sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM, socket.IPPROTO_UDP)
        self.sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        try:
            self.sock.bind(("", MULTICAST_PORT))
        except Exception:
            self.sock.bind((MULTICAST_GROUP, MULTICAST_PORT))

        # Join the Multicast Group on the WiFi interface
        mreq = socket.inet_aton(MULTICAST_GROUP) + socket.inet_aton(lip)
        self.sock.setsockopt(socket.IPPROTO_IP, socket.IP_ADD_MEMBERSHIP, mreq)
        self.sock.setsockopt(socket.IPPROTO_IP, socket.IP_MULTICAST_IF,
                             socket.inet_aton(lip))

    def get_port(self) -> int:
        while True:
            try:
                port = int(input("Enter port: "))
                if 0 < port < 65535:
                    return port
                print("Invalid port")
            except Exception:
                print("Enter numeric value")

    def save_config(self):
        self.config_manager.save_config(
            self.peer_id, self.host, self.registered_nodes,
            self.peers, self.sent, self.recieve)

    def load_config(self):
        config = self.config_manager.load_config()
        if config:
            self.registered_nodes = config.get("registered_nodes", [])
            self.peers = config.get("peers", {})
            self.sent = config.get("sent", [])
            self.recieve = config.get("recieve", [])
        else:
            self.save_config()

    def register_node(self):
        node_name = input("\\nEnter node string: ").strip()
        if node_name == "":
            print("Invalid node")
            return
        with self.lock:
            if node_name not in self.registered_nodes:
                self.registered_nodes.append(node_name)
                self.peers[self.peer_id] = {
                    "host": self.host, "port": self.port,
                    "nodes": self.registered_nodes,
                    "hash": generate_hash(self.peer_id),
                    "time": self.time
                }
                self.save_config()
                print(f"\\nRegistered node: {node_name}")
            else:
                print("Node already exists")

    def view_nodes(self):
        with self.lock:
            print("\\n========== ALL REGISTERED NODES ==========")
            for peer_id, info in self.peers.items():
                print(f"\\nPeer : {peer_id}")
                for node in info["nodes"]:
                    print(f"  - {node}")
            print("\\n==========================================")

    def view_peers(self):
        with self.lock:
            print("\\n=============== PEERS =================")
            for peer_id, info in self.peers.items():
                print(f"\\nPeer ID : {peer_id}")
                print(f"IP      : {info['host']}")
                print(f"Port    : {info['port']}")
                print(f"Hash    : {info['hash'][:20]}...")
            print("=======================================")

    def send(self):
        dest_peer_id = input("\\nEnter destination peer ID (host:port): ").strip()
        if dest_peer_id == "":
            print("Invalid destination peer ID")
            return
        data = input("Enter message to send: ").strip()
        if data == "":
            print("Message cannot be empty")
            return
        with self.lock:
            self.sent.append([data, self.peer_id, dest_peer_id])
            self.save_config()
        print(f"\\nQueued message to {dest_peer_id}: {data}")

    def view_data(self):
        with self.lock:
            print("\\n========== RECEIVED DATA ==========")
            if not self.recieve:
                print("  No data received yet.")
            else:
                for i, entry in enumerate(self.recieve, 1):
                    print(f"  [{i}] Data: {entry[0]} | From: {entry[1]} | To: {entry[2]}")
            print("====================================")

    def command_loop(self):
        while self.running:
            try:
                command = input(f"\\n[{self.port}] Enter command: ")
                if command.lower() == "register":
                    self.register_node()
                elif command.lower() == "view nodes":
                    self.view_nodes()
                elif command.lower() == "view peers":
                    self.view_peers()
                elif command.lower() == "send":
                    self.send()
                elif command.lower() == "view data":
                    self.view_data()
                elif command.lower() == "exit":
                    self.running = False
                    break
            except (KeyboardInterrupt, EOFError):
                self.running = False
                break

    def initialize_self_peer(self):
        with self.lock:
            self.peers[self.peer_id] = {
                "host": self.host, "port": self.port,
                "nodes": self.registered_nodes,
                "hash": generate_hash(self.peer_id),
                "time": self.time
            }
            self.save_config()

    def start(self):
        self.initialize_self_peer()
        print(f"\\nPeer started on {self.peer_id}")
        start_discovery_listener(self)
        start_discovery_sender(self)
        start_health_monitor(self)
        self.command_loop()""")

doc.add_heading('10.3 core/discovery.py — Automatic Peer Discovery', level=2)
add_code_block("""import socket
import json
import time
import threading
from utils.constants import MULTICAST_GROUP, MULTICAST_PORT, DISCOVERY_INTERVAL, BUFFER_SIZE
from utils.hashing import generate_hash

def start_discovery_sender(node) -> threading.Thread:
    def send_loop():
        while node.running:
            message = {
                "type": "DISCOVER",
                "peer_id": node.peer_id,
                "host": node.host,
                "port": node.port,
                "nodes": node.registered_nodes,
                "hash": generate_hash(node.peer_id),
                "sent": node.sent
            }
            try:
                node.sock.sendto(
                    json.dumps(message).encode(),
                    (MULTICAST_GROUP, MULTICAST_PORT))
            except Exception:
                pass
            time.sleep(DISCOVERY_INTERVAL)
    thread = threading.Thread(target=send_loop, daemon=True)
    thread.start()
    return thread

def start_discovery_listener(node) -> threading.Thread:
    def listen_loop():
        while node.running:
            try:
                data, addr = node.sock.recvfrom(BUFFER_SIZE)
                message = json.loads(data.decode())
                if message["peer_id"] == node.peer_id:
                    continue
                for i in message["sent"]:
                    if i[2] == node.peer_id:
                        with node.lock:
                            if i not in node.recieve:
                                node.recieve.append(i)
                                node.save_config()
                if message["type"] == "DISCOVER":
                    with node.lock:
                        node.peers[message["peer_id"]] = {
                            "host": message["host"],
                            "port": message["port"],
                            "nodes": message["nodes"],
                            "hash": message["hash"],
                            "time": node.time,
                            "recieve": node.recieve,
                        }
                        node.save_config()
            except Exception:
                pass
    thread = threading.Thread(target=listen_loop, daemon=True)
    thread.start()
    return thread""")

doc.add_heading('10.4 core/health_monitor.py — Health Monitoring', level=2)
add_code_block("""import time
import threading
from utils.constants import PEER_TIMEOUT

def start_health_monitor(node) -> threading.Thread:
    def monitor_loop():
        while node.running:
            remove_peers = []
            with node.lock:
                for peer_id, info in list(node.peers.items()):
                    if peer_id == node.peer_id:
                        continue
                    info["time"] += 2
                    if info["time"] > PEER_TIMEOUT:
                        remove_peers.append(peer_id)
                if remove_peers:
                    for peer_id in remove_peers:
                        del node.peers[peer_id]
                        print(f"\\nRemoved inactive peer: {peer_id}")
                    node.save_config()
            time.sleep(2)
    thread = threading.Thread(target=monitor_loop, daemon=True)
    thread.start()
    return thread""")

doc.add_heading('10.5 core/config_manager.py — State Persistence', level=2)
add_code_block("""import os
import json

class ConfigManager:
    def __init__(self, port):
        self.port = port
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.storage_dir = os.path.join(base_dir, "storage", "configs")
        self.config_file = os.path.join(self.storage_dir, f"{port}.json")
        os.makedirs(self.storage_dir, exist_ok=True)

    def save_config(self, peer_id, host, registered_nodes, peers, sent, recieve):
        data = {
            "peer_id": peer_id, "host": host,
            "port": self.port, "registered_nodes": registered_nodes,
            "peers": peers, "sent": sent, "recieve": recieve
        }
        with open(self.config_file, "w") as f:
            json.dump(data, f, indent=4)

    def load_config(self):
        if os.path.exists(self.config_file):
            with open(self.config_file, "r") as f:
                return json.load(f)
        return None""")

doc.add_heading('10.6 utils/constants.py', level=2)
add_code_block("""# Network constants for the P2P multicast system

MULTICAST_GROUP = "224.1.1.1"
MULTICAST_PORT = 5007

BUFFER_SIZE = 4096

DISCOVERY_INTERVAL = 5
PEER_TIMEOUT = 10""")

doc.add_heading('10.7 utils/hashing.py', level=2)
add_code_block("""import hashlib

def generate_hash(value: str) -> str:
    \"\"\"Generates a SHA-256 hash for a given string value.\"\"\"
    return hashlib.sha256(value.encode()).hexdigest()""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 11: FAULT DIAGNOSTIC SYSTEM
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Fault Diagnostic System', level=1)

add_body(
    'This chapter presents a comprehensive description of the Fault Diagnostic System (FDS) in the '
    'context of this project. It describes the role of fault diagnosis in integrated sonar suites, '
    'the limitations of current monolithic FDS architectures, and how the PTP-DDS approach implemented '
    'in this project addresses those limitations by distributing fault diagnosis intelligence across '
    'multiple peer nodes.'
)

doc.add_heading('11.1 Overview of Fault Diagnosis in Sonar Suites', level=2)
add_body(
    'In the context of this project — "Implementing a Lightweight PTP-DDS for Building a Resilient '
    'Fault Diagnosis Information Sharing Network in Integrated Sonar Suites" — the Fault Diagnosis '
    'System (FDS) is the subsystem responsible for continuously monitoring the health of all software, '
    'hardware, and networking components within an integrated sonar suite, detecting failures, and '
    'sharing diagnostic information so that corrective action can be taken.'
)
add_body(
    'Integrated sonar suites are complex, distributed, real-time computing environments comprising '
    'multiple interconnected processing nodes — beam-formers, signal processors, target detectors, '
    'display consoles, and recorders. Each of these subsystems can fail independently due to hardware '
    'faults, software crashes, network disconnections, or environmental factors. The FDS must continuously '
    'monitor all these subsystems, detect any failures in real time, isolate the faulty component, and '
    'share diagnostic information across the system so that operators and automated recovery mechanisms '
    'can take appropriate corrective action.'
)
add_body(
    'Fault diagnosis in sonar systems encompasses three core capabilities: (1) Fault Detection — the '
    'ability to identify that a component has failed or is no longer functioning correctly; (2) Fault '
    'Isolation — the ability to pinpoint the specific failed component without affecting the rest of '
    'the system; and (3) Fault Recovery — the ability to restore normal operation after a fault has '
    'been detected and isolated, either through automatic mechanisms or operator intervention.'
)

doc.add_heading('11.2 Limitations of Current Monolithic FDS', level=2)
add_body(
    'Current generation Fault Diagnosis Systems in sonar suites follow a centralized, monolithic design '
    'where a single FDS application runs on a dedicated processor. This application collects health status '
    'from all subsystems via the ship-board LAN, performs diagnostic analysis, and reports faults to operators. '
    'While this approach is straightforward to implement, it introduces several critical limitations:'
)

add_table(
    ['Limitation', 'Description', 'Operational Impact'],
    [
        ['Single Point of Failure', 'The FDS application runs on a single dedicated processor. If this processor crashes, hangs, or becomes network-unreachable, fault diagnosis capability for the entire sonar suite is lost.', 'The system designed to detect faults is itself vulnerable to a single fault — operationally unacceptable for mission-critical naval systems.'],
        ['Static Configuration', 'Every monitored subsystem must be manually configured with its network address and monitoring parameters in the FDS application.', 'Adding, removing, or replacing subsystems (e.g., during hardware upgrades or field maintenance) requires manual FDS reconfiguration — time-consuming and error-prone.'],
        ['No Dynamic Discovery', 'The FDS has no mechanism to autonomously detect the presence of new subsystems or the absence of failed ones without explicit configuration.', 'Failed nodes may go undetected until manual inspection; new nodes require manual registration.'],
        ['Scalability Bottleneck', 'As sonar suites grow in complexity with more processing nodes and sensors, a single centralized FDS struggles to handle increasing volumes of health status data.', 'The centralized architecture creates a processing and communication bottleneck that limits system scalability.'],
        ['Tight Coupling', 'The FDS application is tightly coupled to the specific set of subsystems it monitors, requiring code or configuration changes when the monitored topology changes.', 'Reduced flexibility and increased maintenance burden during system evolution.'],
    ]
)

add_body(
    'These limitations are fundamentally architectural — they cannot be resolved by improving the '
    'monolithic FDS application itself. A fundamentally different approach is needed: one that distributes '
    'fault diagnosis capability across all nodes in the system, eliminating the single point of failure '
    'and enabling automatic, zero-configuration operation.'
)

doc.add_heading('11.3 PTP-DDS Approach to Fault Diagnosis', level=2)
add_body(
    'The PTP-DDS system implemented in this project replaces the centralized monolithic FDS with a '
    'decentralized, peer-to-peer fault diagnosis network. Instead of a single FDS processor, every '
    'node in the sonar suite becomes a fault diagnosis participant. The Fault Diagnostic System in '
    'this project is not a single centralized application — it is the collective behaviour of all peer '
    'nodes working together through the PTP-DDS middleware.'
)
add_body(
    'The key mechanisms that enable distributed fault diagnosis in the PTP-DDS system are:'
)

add_body('Mechanism 1: Automatic Peer Discovery (core/discovery.py)')
add_body(
    'Nodes join a well-known UDP multicast group (224.1.1.1:5007) and broadcast DISCOVER messages every '
    '5 seconds containing their identity (peer_id, host, port, registered nodes, hash) and queued messages. '
    'Any new node joining the LAN is instantly detected by all existing nodes without any manual '
    'configuration. This mechanism directly mirrors the DDS Simple Participant Discovery Protocol (SPDP) '
    'and serves as the foundation of the fault diagnosis network — a node must be discovered before it '
    'can be monitored.'
)

add_body('Mechanism 2: Heartbeat-Based Health Monitoring (core/health_monitor.py)')
add_body(
    'A dedicated background thread — the Health Monitor — continuously checks when each remote peer '
    'was last heard from. Every 2 seconds, it increments a logical timer for each remote peer. When '
    'a DISCOVER message is received from a peer, that peer\'s timer is reset to zero, effectively '
    'serving as a heartbeat acknowledgement. If a peer\'s timer exceeds the PEER_TIMEOUT threshold '
    '(10 seconds — indicating two missed heartbeats), the peer is automatically declared dead and '
    'pruned from the active peer list. This is the core fault detection mechanism of the system, '
    'directly analogous to the DDS Liveliness QoS policy.'
)

add_body('Mechanism 3: Distributed Fault Information Sharing (core/node.py — send/view_data)')
add_body(
    'Nodes can exchange directed messages (fault reports, status updates, diagnostic data) peer-to-peer. '
    'Sent messages are piggybacked onto the periodic DISCOVER broadcast and delivered to the destination '
    'peer, enabling any node to share diagnostic information with any other node without a central broker. '
    'This enables distributed fault awareness — every node can inform every other node about faults it '
    'has detected in its local subsystem.'
)

add_body('Mechanism 4: State Persistence and Recovery (core/config_manager.py)')
add_body(
    'Each node persists its complete state — peer table, registered nodes, sent messages, and received '
    'messages — to a JSON file on disk (storage/configs/<port>.json). If a node crashes and restarts, '
    'it automatically recovers its previous state and re-joins the fault diagnosis network. This provides '
    'durability analogous to the DDS Durability QoS (TRANSIENT_LOCAL) policy, ensuring that fault '
    'diagnosis history is not lost during node restarts.'
)

add_body('Mechanism 5: Node/Topic Registration (core/node.py — register_node)')
add_body(
    'Nodes register named channels representing the subsystems they monitor (analogous to DDS Topics, '
    'e.g., "BeamFormer-1", "SignalProcessor-2"). These registrations are disseminated to all peers '
    'as part of the DISCOVER broadcast, so every node in the fault diagnosis network knows what every '
    'other node is monitoring. This enables comprehensive system-wide fault awareness.'
)

doc.add_heading('11.4 Fault Diagnosis Lifecycle', level=2)
add_body(
    'The fault diagnosis lifecycle in the PTP-DDS system operates as a continuous, automated cycle '
    'across all peer nodes. The following describes the complete lifecycle from node startup through '
    'fault detection, isolation, and recovery:'
)

add_code_block("""Node Starts → Joins Multicast Group → Broadcasts DISCOVER every 5s
                                          |
    +-----------------------------------------+
    v
Other Nodes Receive DISCOVER → Update Peer Table → Reset Peer Timer
                                                        |
                                               Timer exceeds 10s?
                                              /                  \\
                                           No                    Yes
                                            |                     |
                                      Peer is alive        +-------------+
                                                           | FAULT       |
                                                           | DETECTED    |
                                                           | Remove peer |
                                                           | Save state  |
                                                           | Log removal |
                                                           +-------------+""")

add_body('Stage 1: Discovery and Registration')
add_body(
    'When a node starts, it creates a UDP multicast socket, joins the multicast group, loads any '
    'previously saved state from disk, and begins broadcasting DISCOVER messages every 5 seconds. '
    'Simultaneously, it listens for DISCOVER messages from other nodes. Within one discovery interval '
    '(5 seconds), the node has discovered all active peers on the LAN and has been discovered by them. '
    'No manual configuration is required — this is zero-configuration, plug-and-play fault diagnosis.'
)

add_body('Stage 2: Continuous Health Monitoring')
add_body(
    'Once peers are discovered, the Health Monitor thread begins tracking their liveliness. Each '
    'received DISCOVER message resets the sending peer\'s timer to zero. The monitor increments all '
    'remote peer timers by 2 every 2 seconds. As long as a peer continues sending DISCOVER messages, '
    'its timer stays near zero, confirming it is alive and healthy.'
)

add_body('Stage 3: Fault Detection')
add_body(
    'If a peer stops sending DISCOVER messages (due to a crash, network disconnection, or shutdown), '
    'its timer continues to increment without being reset. After 10 seconds (the PEER_TIMEOUT threshold), '
    'the Health Monitor declares the peer as faulty. This provides deterministic, bounded fault detection '
    'time — a fault is guaranteed to be detected within PEER_TIMEOUT + one monitor cycle (10–12 seconds).'
)

add_body('Stage 4: Fault Isolation and Notification')
add_body(
    'The faulty peer is removed from the active peer table, the updated state is saved to disk, and '
    'a diagnostic message ("Removed inactive peer: <peer_id>") is logged. Crucially, only the specific '
    'faulty peer is affected — all other peers and their data channels continue operating normally. '
    'This per-peer granularity of fault isolation is a direct benefit of the decentralized P2P '
    'architecture, where there is no central broker whose failure could cascade across the system.'
)

add_body('Stage 5: Fault Recovery')
add_body(
    'When a failed node is repaired and restarted, it automatically re-joins the multicast group and '
    'begins broadcasting DISCOVER messages. All surviving nodes detect the new DISCOVER messages and '
    're-register the recovered peer within 5 seconds — with zero manual intervention. The restarted '
    'node also loads its previous state from disk, restoring its registered nodes, peer history, and '
    'message logs. This enables seamless fault recovery, critical for sonar systems where field '
    'replacement of processing nodes must be operationally transparent.'
)

doc.add_heading('11.5 Advantages Over Monolithic FDS', level=2)
add_body(
    'The PTP-DDS approach to fault diagnosis offers significant advantages over the traditional '
    'monolithic FDS architecture. The following table provides a detailed comparison:'
)

add_table(
    ['Feature', 'Monolithic FDS', 'PTP-DDS Fault Diagnosis Network'],
    [
        ['Single Point of Failure', 'FDS processor is a single point of failure — its failure causes total loss of fault diagnosis capability', 'No single point of failure — fault diagnosis is distributed across all peer nodes; any surviving node continues monitoring'],
        ['Discovery', 'Manual IP configuration — every monitored subsystem must be explicitly configured in the FDS application', 'Zero-configuration UDP multicast — new nodes are automatically discovered within 5 seconds'],
        ['Scalability', 'Central bottleneck — all health data converges on a single FDS processor', 'Linear scalability — each new node is self-managing and independently monitors all peers'],
        ['Fault Detection', 'Only the centralized FDS can detect faults in other subsystems', 'Every node independently monitors every other node via heartbeat-based health checking'],
        ['Fault Isolation', 'FDS failure affects all fault diagnosis capability', 'Per-peer fault isolation — only the faulty peer is removed; no cascading failures'],
        ['Recovery', 'Manual restart and reconfiguration required after FDS failure', 'Automatic state recovery from disk and zero-configuration re-discovery on restart'],
        ['Graceful Degradation', 'Total loss of diagnostics if FDS fails', 'Remaining nodes continue full fault diagnosis operation; system degrades gracefully'],
        ['Maintenance', 'FDS configuration must be manually updated when subsystems are added or removed', 'Plug-and-play — nodes self-register and self-deregister automatically'],
        ['DDS Alignment', 'No alignment with DDS middleware standards', 'Directly implements DDS concepts: SPDP discovery, Liveliness QoS, Durability QoS'],
    ]
)

add_body(
    'In summary, the Fault Diagnostic System in this project is the collective behaviour of all peer '
    'nodes working together through the PTP-DDS middleware. Each node independently discovers peers, '
    'monitors their health via heartbeats, detects failures through timeout-based pruning, shares '
    'diagnostic information via peer-to-peer messaging, and persists state for crash recovery. This '
    'decentralized approach ensures that no single node failure can take down the fault diagnosis '
    'capability of the entire sonar suite — directly addressing the most critical limitation of '
    'current monolithic FDS architectures.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 12: RELEVANCE TO SONAR SYSTEMS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Relevance to Sonar Systems', level=1)

add_body(
    'The PTP-DDS prototype directly addresses the core problem of building a resilient fault diagnosis '
    'information sharing network for integrated sonar suites. This chapter maps the prototype capabilities '
    'to actual sonar system requirements and demonstrates how the PTP-DDS approach eliminates the '
    'single point of failure inherent in current monolithic FDS architectures.'
)

doc.add_heading('12.1 Mapping to Sonar FDS Architecture', level=2)
add_table(
    ['Current Monolithic FDS', 'PTP-DDS Fault Diagnosis Network'],
    [
        ['Single FDS application on one processor', 'Fault diagnosis distributed across all PTP-DDS peer nodes'],
        ['FDS is a single point of failure', 'No single point of failure — any surviving node continues monitoring'],
        ['Manual configuration of monitored subsystems', 'Automatic zero-configuration discovery of all peers'],
        ['Static fault monitoring topology', 'Dynamic topology — nodes join/leave automatically'],
        ['Centralized fault database', 'Distributed fault information sharing via P2P messaging'],
        ['FDS failure = total loss of diagnostics', 'Graceful degradation — remaining nodes continue diagnosis'],
    ]
)

add_body('Mapping to sonar processing nodes:')
add_table(
    ['Sonar Subsystem', 'PTP-DDS Analogy'],
    [
        ['Beam Former Node', 'PTP-DDS peer publishing "beam_health" diagnostic topic'],
        ['Signal Processor Node', 'PTP-DDS peer sharing processing health and fault status'],
        ['Display Console', 'PTP-DDS peer subscribing to fault status from all subsystems'],
        ['Recorder', 'PTP-DDS peer with persistent fault history (state persistence)'],
        ['Redundant FDS Node', 'Hot-standby peer that auto-discovers and continues diagnosis on failure'],
    ]
)

doc.add_heading('12.2 Key Benefits for Sonar FDS', level=2)
add_bullet('Resilience — Eliminates the single point of failure of monolithic FDS. Fault diagnosis capability survives individual node failures.')
add_bullet('Easy Scalability — New subsystems are automatically discovered and integrated into the fault diagnosis network without configuration changes.')
add_bullet('Plug-and-Play Deployment — Replacement nodes or additional monitoring points can be added to the sonar suite without any manual configuration.')
add_bullet('Distributed Fault Awareness — All nodes share fault information, enabling any node to act on diagnostic data.')
add_bullet('Deterministic Health Monitoring — The logical timer-based health check provides deterministic, reproducible fault detection in embedded real-time environments.')
add_bullet('State Recovery — Nodes recover their previous fault diagnosis state after restarts, minimizing diagnostic gaps during maintenance.')

doc.add_heading('12.3 Limitations and Considerations', level=2)
add_table(
    ['Limitation', 'Mitigation in Production DDS'],
    [
        ['JSON encoding overhead', 'DDS uses CDR binary serialization for zero-copy performance'],
        ['Python GIL limits true parallelism', 'Production DDS uses C/C++ with zero-copy shared memory'],
        ['No QoS policy negotiation', 'Full DDS provides 22+ QoS policies with automatic matching'],
        ['UDP multicast limited to LAN', 'DDS supports UDP unicast, TCP, and shared-memory transports'],
        ['No data filtering', 'DDS supports content-filtered topics and time-based filters'],
        ['No encryption/authentication', 'DDS Security specification provides end-to-end encryption'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 13: CONCLUSION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Conclusion', level=1)
add_body(
    'This internship at DRDO-NPOL, Thrikkakara, addressed a critical limitation in current sonar system '
    'architectures: the single point of failure inherent in monolithic Fault Diagnosis Systems. '
    'A lightweight Peer-to-Peer Data Distribution Service (PTP-DDS) was designed and implemented '
    'to build a resilient fault diagnosis information sharing network for integrated sonar suites. '
    'The key outcomes of this project are:'
)
add_bullet(
    'Comprehensive understanding of DDS architecture — The study covered the DCPS and RTPS layers, '
    'QoS policies, automatic discovery protocols (SPDP/SEDP), and the data-centric publish-subscribe paradigm.'
)
add_bullet(
    'Rigorous analysis of P2P network designs — Four major approaches for realising peer-to-peer networks '
    'and data distribution services (centralized registry, gossip protocol, DHT, UDP multicast) were studied '
    'and evaluated for suitability to build a fault diagnosis information sharing network. The UDP multicast '
    'approach was selected as optimal for LAN-scoped sonar suites due to its zero-configuration, instant '
    'discovery, full decentralization, and direct alignment with DDS SPDP.'
)
add_bullet(
    'Successful PTP-DDS prototype library — A fully functional PTP-DDS library was developed in Python '
    'implementing core capabilities: automatic peer discovery via UDP multicast, heartbeat-based fault '
    'detection with deterministic timeout bounds, fault isolation at the individual peer level, automatic '
    'fault recovery through re-discovery and state persistence, and thread-safe concurrent operation. '
    'The library eliminates the single point of failure of monolithic FDS architectures by distributing '
    'fault diagnosis across all peer nodes.'
)
add_bullet(
    'Validation through testing and fault injection — The system was validated with 7 unit tests covering '
    'all core components, all passing successfully. Fault injection scenarios demonstrated reliable fault '
    'detection (within 12 seconds), fault isolation (no cascading failures), fault recovery '
    '(zero-configuration re-discovery), and system resilience under cascading failures.'
)
add_bullet(
    'Feasibility assessment — The PTP-DDS approach directly addresses the limitations of current '
    'monolithic FDS: decentralization eliminates the single point of failure, automatic discovery enables '
    'plug-and-play scalability, heartbeat-based monitoring provides deterministic fault detection, and '
    'the data-centric model enables distributed fault information sharing across all nodes.'
)
add_body(
    'The PTP-DDS prototype serves as a foundational library that validates the core design for building '
    'a resilient fault diagnosis information sharing network. With production DDS implementations '
    '(such as RTI Connext DDS or OpenDDS), these concepts can be directly applied to realize '
    'next-generation integrated sonar suites with distributed, resilient fault diagnosis capabilities.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 14: FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('14. Future Scope', level=1)
add_bullet('Integration with production DDS — Port the PTP-DDS library to a production DDS implementation (e.g., RTI Connext DDS, OpenDDS) with full QoS support for deployment in actual sonar suites.')
add_bullet('Sonar FDS data simulation — Generate synthetic fault diagnosis data (health status, error codes, subsystem states) and distribute them through the PTP-DDS network.')
add_bullet('QoS policy evaluation — Systematically benchmark different QoS configurations (RELIABLE vs BEST_EFFORT, various history depths) for fault diagnosis data streams.')
add_bullet('Multi-topic fault taxonomy — Implement separate DDS topics for different fault categories (hardware faults, software faults, network faults, environmental faults) with appropriate QoS per category.')
add_bullet('Performance benchmarking — Measure latency, throughput, and jitter of fault information dissemination under realistic sonar workloads.')
add_bullet('Cross-platform deployment — Deploy PTP-DDS nodes across heterogeneous hardware (x86 servers, ARM-based embedded boards) representative of actual sonar system configurations.')
add_bullet('Security integration — Evaluate DDS Security specification for encrypted, authenticated fault diagnosis data distribution.')
add_bullet('GUI Dashboard — Develop a web-based or desktop GUI for real-time visualization of peer topology, fault status, and system health across the sonar suite.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CHAPTER 15: REFERENCES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('15. References', level=1)

references = [
    'Object Management Group, "Data Distribution Service for Real-Time Systems (DDS), Version 1.4," OMG Document formal/2015-04-10, 2015.',
    'Object Management Group, "The Real-Time Publish-Subscribe Wire Protocol DDS Interoperability Wire Protocol (DDSI-RTPS), Version 2.3," OMG Document formal/2019-04-03, 2019.',
    'G. Pardo-Castellote, "OMG Data-Distribution Service: Architectural Overview," 23rd International Conference on Distributed Computing Systems Workshops, Providence, RI, 2003, pp. 200-206.',
    'Real-Time Innovations, "RTI Connext DDS — The Connectivity Framework for IIoT," RTI Documentation, 2024.',
    'Object Computing Inc., "OpenDDS Developer\'s Guide," Version 3.27, 2024.',
    'A. Corsaro, "Data-Centric Middleware for Autonomous and Intelligent Systems," IEEE International Conference on Autonomous Robot Systems and Competitions, 2019.',
    'J. M. López-Higuera, "Sonar Signal Processing," in Handbook of Optical Fibre Sensing Technology, John Wiley & Sons, 2002.',
    'R. J. Urick, Principles of Underwater Sound, 3rd ed., Peninsula Publishing, 1983.',
    'Python Software Foundation, "socket — Low-level networking interface," Python 3.12 Documentation, 2024.',
    'Python Software Foundation, "threading — Thread-based parallelism," Python 3.12 Documentation, 2024.',
    'eProsima, "Fast DDS Documentation," Version 2.14, 2024.',
    'I. Stoica et al., "Chord: A Scalable Peer-to-peer Lookup Protocol for Internet Applications," IEEE/ACM Transactions on Networking, vol. 11, no. 1, pp. 17-32, 2003.',
    'P. Maymounkov and D. Mazières, "Kademlia: A Peer-to-Peer Information System Based on the XOR Metric," Proc. IPTPS, 2002.',
    'A. Demers et al., "Epidemic Algorithms for Replicated Database Maintenance," Proc. 6th ACM PODC, 1987.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(f"[{i}]  {ref}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

# ── Save ──
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Internship_Report_DDS_Sonar_System.docx"
)
doc.save(output_path)
print(f"\nReport saved to: {output_path}")
